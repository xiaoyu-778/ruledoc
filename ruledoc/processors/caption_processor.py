"""
CaptionProcessor - 题注处理器

职责:
- 识别图表题注
- 生成正确的编号 (章节号-序号)
- 自动检测题注所属章节

执行顺序:
CaptionProcessor 是第三个执行的处理器，在 StyleProcessor 和 HeadingProcessor 之后。
注意: CaptionProcessor 会自动检测题注所属章节，不依赖 ctx.chapter_num 的最终值。
"""

import re
from typing import TYPE_CHECKING, Any, Dict, List, Optional

from ruledoc.processors.base import PostProcessor

if TYPE_CHECKING:
    from ruledoc.context import ProcessingContext


class CaptionProcessor(PostProcessor):
    """
    题注处理器

    处理文档中的图表题注:
    - 识别题注段落 (基于样式或文本模式)
    - 自动检测题注所属章节
    - 生成正确的编号格式 "图 X-Y 标题" 或 "表 X-Y 标题"

    设计说明:
    CaptionProcessor 会先扫描文档确定每个段落所属的章节，
    然后处理题注时使用正确的章节号。这解决了处理器独立遍历导致的
    章节号不同步问题。

    Attributes:
        fig_patterns: 图题注匹配模式列表
        tab_patterns: 表题注匹配模式列表
    """

    DEFAULT_FIG_PATTERNS = [
        re.compile(r"^(图\s*)(\d+[-‐–—]\d+)?\s*(.+)$"),
        re.compile(r"^(Figure\s*)(\d+[-‐–—]\d+)?\s*(.+)$", re.IGNORECASE),
        re.compile(r"^(Fig\.?\s*)(\d+[-‐–—]\d+)?\s*(.+)$", re.IGNORECASE),
    ]

    DEFAULT_TAB_PATTERNS = [
        re.compile(r"^(表\s*)(\d+[-‐–—]\d+)?\s*(.+)$"),
        re.compile(r"(Table\s*)(\d+[-‐–—]\d+)?\s*(.+)$", re.IGNORECASE),
        re.compile(r"(Tab\.?\s*)(\d+[-‐–—]\d+)?\s*(.+)$", re.IGNORECASE),
    ]

    CAPTION_SIMPLE_NUM_PATTERN = re.compile(r"^(图|表|Figure|Table|Fig\.?|Tab\.?)\s*\d+[.\s]\s*")

    CAPTION_STYLES = {
        "caption",
        "figure",
        "table",
        "figure caption",
        "table caption",
        "题注",
        "图表题注",
    }

    HEADING_PATTERN = re.compile(r"^heading\s*(\d+)$", re.IGNORECASE)

    NON_CHAPTER_HEADINGS = {
        "摘要",
        "abstract",
        "关键词",
        "keywords",
        "参考文献",
        "references",
        "目录",
        "contents",
        "致谢",
        "acknowledgements",
        "附录",
        "appendix",
        "结论",
        "conclusion",
        "引言",
        "introduction",
    }

    def __init__(
        self,
        fig_patterns: Optional[List[re.Pattern]] = None,
        tab_patterns: Optional[List[re.Pattern]] = None,
        use_style_detection: bool = True,
        use_seq_field: bool = False,
        non_chapter_headings: Optional[List[str]] = None,
    ):
        """
        初始化题注处理器

        Args:
            fig_patterns: 自定义图题注匹配模式列表
            tab_patterns: 自定义表题注匹配模式列表
            use_style_detection: 是否启用样式检测
            use_seq_field: 是否使用SEQ域生成自动编号（默认False，MS Word兼容性更好）
            non_chapter_headings: 不参与章节计数的标题关键词列表
        """
        self._fig_patterns = fig_patterns or self.DEFAULT_FIG_PATTERNS
        self._tab_patterns = tab_patterns or self.DEFAULT_TAB_PATTERNS
        self._use_style_detection = use_style_detection
        self._use_seq_field = use_seq_field
        self._non_chapter_headings = non_chapter_headings or list(self.NON_CHAPTER_HEADINGS)

        self._processed_figs = 0
        self._processed_tabs = 0
        self._skipped_captions = 0

        self._chapter_counters: Dict[int, Dict[str, int]] = {}

    def process(self, ctx: "ProcessingContext") -> None:
        """
        执行题注处理

        先扫描文档确定每个段落所属的章节，然后处理题注。

        Args:
            ctx: 处理上下文
        """
        self._processed_figs = 0
        self._processed_tabs = 0
        self._skipped_captions = 0
        self._chapter_counters = {}

        paragraph_chapters = self._scan_chapters(ctx.doc)

        for i, para in enumerate(ctx.doc.paragraphs):
            if self._is_caption(para):
                chapter_num = paragraph_chapters.get(i, 1)
                self._process_caption(para, ctx, chapter_num)

        ctx.add_warning(
            f"CaptionProcessor: 处理了 {self._processed_figs} 个图题注, "
            f"{self._processed_tabs} 个表题注, 跳过 {self._skipped_captions} 个"
        )

    def _scan_chapters(self, doc) -> Dict[int, int]:
        """
        扫描文档，确定每个段落所属的章节号

        仅将真正的章节一级标题计入章节号，
        跳过摘要、参考文献、致谢等特殊标题和论文标题（第一个H1）。

        Args:
            doc: 文档对象

        Returns:
            字典: {段落索引: 章节号}
        """
        paragraph_chapters: Dict[int, int] = {}
        current_chapter = 0
        first_h1_seen = False

        for i, para in enumerate(doc.paragraphs):
            if self._is_heading_level_1(para):
                if not first_h1_seen:
                    first_h1_seen = True
                    paragraph_chapters[i] = current_chapter if current_chapter > 0 else 1
                    continue

                text = para.text.strip().lower().replace(" ", "")
                is_non_chapter = any(keyword in text for keyword in self._non_chapter_headings)
                if not is_non_chapter:
                    current_chapter += 1

            paragraph_chapters[i] = current_chapter if current_chapter > 0 else 1

        return paragraph_chapters

    def _is_heading_level_1(self, para) -> bool:
        """
        检测段落是否为一级标题

        Args:
            para: 段落对象

        Returns:
            是否为一级标题
        """
        if not para.style:
            return False

        style_name = para.style.name.lower()
        match = self.HEADING_PATTERN.match(style_name)
        if match and match.group(1) == "1":
            return True

        return "heading 1" in style_name or style_name == "heading1"

    def _is_caption(self, para) -> bool:
        """
        检测段落是否为题注

        通过样式或文本模式识别题注。

        Args:
            para: 段落对象

        Returns:
            是否为题注
        """
        text = para.text.strip()
        if not text:
            return False

        if self._use_style_detection and para.style:
            style_name = para.style.name.lower()
            if style_name in self.CAPTION_STYLES:
                return True

        for pattern in self._fig_patterns:
            if pattern.match(text):
                return True

        for pattern in self._tab_patterns:
            if pattern.match(text):
                return True

        return False

    def _process_caption(self, para, ctx: "ProcessingContext", chapter_num: int) -> None:
        """
        处理单个题注段落

        Args:
            para: 段落对象
            ctx: 处理上下文
            chapter_num: 该段落所属的章节号
        """
        text = para.text.strip()

        for pattern in self._fig_patterns:
            match = pattern.match(text)
            if match:
                self._process_figure_caption(para, match, ctx, chapter_num)
                return

        for pattern in self._tab_patterns:
            match = pattern.match(text)
            if match:
                self._process_table_caption(para, match, ctx, chapter_num)
                return

        self._skipped_captions += 1
        ctx.add_warning(f"CaptionProcessor: 无法识别的题注格式: {text[:30]}...")

    def _get_next_fig_number(self, chapter_num: int) -> str:
        """
        获取指定章节的下一个图编号

        Args:
            chapter_num: 章节号

        Returns:
            格式为 "章节号-图号" 的字符串
        """
        if chapter_num not in self._chapter_counters:
            self._chapter_counters[chapter_num] = {"fig": 0, "tab": 0}

        self._chapter_counters[chapter_num]["fig"] += 1
        return f"{chapter_num}-{self._chapter_counters[chapter_num]['fig']}"

    def _get_next_tab_number(self, chapter_num: int) -> str:
        """
        获取指定章节的下一个表编号

        Args:
            chapter_num: 章节号

        Returns:
            格式为 "章节号-表号" 的字符串
        """
        if chapter_num not in self._chapter_counters:
            self._chapter_counters[chapter_num] = {"fig": 0, "tab": 0}

        self._chapter_counters[chapter_num]["tab"] += 1
        return f"{chapter_num}-{self._chapter_counters[chapter_num]['tab']}"

    def _process_figure_caption(
        self, para, match, ctx: "ProcessingContext", chapter_num: int
    ) -> None:
        """
        处理图题注

        Args:
            para: 段落对象
            match: 正则匹配对象
            ctx: 处理上下文
            chapter_num: 章节号
        """
        groups = match.groups()
        raw_caption = groups[2] if len(groups) > 2 else ""

        caption_text = self._strip_caption_number(raw_caption)

        if not caption_text:
            self._skipped_captions += 1
            return

        font_cfg = self._get_font_config(ctx)

        item_num = self._get_next_fig_item_number(chapter_num)
        number = f"{chapter_num}-{item_num}"

        if self._use_seq_field:
            self._format_caption_with_seq_field(
                para,
                "图",
                chapter_num,
                item_num,
                caption_text.strip(),
                chinese_font=font_cfg["chinese_font"],
                english_font=font_cfg["english_font"],
                caption_size=font_cfg["caption_size"],
            )
        else:
            self._format_caption_static(
                para,
                "图",
                number,
                caption_text.strip(),
                chinese_font=font_cfg["chinese_font"],
                english_font=font_cfg["english_font"],
                caption_size=font_cfg["caption_size"],
            )

        self._processed_figs += 1
        ctx.add_warning(f"CaptionProcessor: 图 {number} - {caption_text[:20]}...")

    def _process_table_caption(
        self, para, match, ctx: "ProcessingContext", chapter_num: int
    ) -> None:
        """
        处理表题注

        Args:
            para: 段落对象
            match: 正则匹配对象
            ctx: 处理上下文
            chapter_num: 章节号
        """
        groups = match.groups()
        raw_caption = groups[2] if len(groups) > 2 else ""

        caption_text = self._strip_caption_number(raw_caption)

        if not caption_text:
            self._skipped_captions += 1
            return

        font_cfg = self._get_font_config(ctx)

        item_num = self._get_next_tab_item_number(chapter_num)
        number = f"{chapter_num}-{item_num}"

        if self._use_seq_field:
            self._format_caption_with_seq_field(
                para,
                "表",
                chapter_num,
                item_num,
                caption_text.strip(),
                chinese_font=font_cfg["chinese_font"],
                english_font=font_cfg["english_font"],
                caption_size=font_cfg["caption_size"],
            )
        else:
            self._format_caption_static(
                para,
                "表",
                number,
                caption_text.strip(),
                chinese_font=font_cfg["chinese_font"],
                english_font=font_cfg["english_font"],
                caption_size=font_cfg["caption_size"],
            )

        self._processed_tabs += 1
        ctx.add_warning(f"CaptionProcessor: 表 {number} - {caption_text[:20]}...")

    def _get_font_config(self, ctx: "ProcessingContext") -> Dict[str, Any]:
        """
        从处理上下文的规则配置中提取字体设置

        优先使用规则配置中的字体和字号，
        若规则未提供则回退到学术规范默认值（宋体 + Times New Roman + 10.5pt）。

        Args:
            ctx: 处理上下文

        Returns:
            字体配置字典 {
                'chinese_font': 中文字体名,
                'english_font': 英文字体名,
                'caption_size': 题注字号(pt)
            }
        """
        defaults = {
            "chinese_font": "宋体",
            "english_font": "Times New Roman",
            "caption_size": 10.5,
        }

        try:
            rule = getattr(ctx, "rule", None)
            if rule is None:
                return defaults

            font_settings = getattr(rule, "FONT_SETTINGS", None)
            font_sizes = getattr(rule, "FONT_SIZES", None)

            config = dict(defaults)
            if font_settings and isinstance(font_settings, dict):
                if "chinese_font" in font_settings:
                    config["chinese_font"] = font_settings["chinese_font"]
                if "english_font" in font_settings:
                    config["english_font"] = font_settings["english_font"]
            if font_sizes and isinstance(font_sizes, dict) and "caption" in font_sizes:
                config["caption_size"] = font_sizes["caption"]

            return config
        except Exception:
            return defaults

    def _strip_caption_number(self, text: str) -> str:
        """
        去除题注文本中原有的简单数字编号

        处理 "1展示了系统架构图" → "展示了系统架构图"
             "1 系统架构图" → "系统架构图"
             "展示了系统架构图" → "展示了系统架构图"

        Args:
            text: 原始题注文本

        Returns:
            去除编号后的纯标题文本
        """
        text = text.strip()
        stripped = self.CAPTION_SIMPLE_NUM_PATTERN.sub("", text)
        return stripped.strip()

    def _get_next_fig_item_number(self, chapter_num: int) -> int:
        """获取指定章节的下一个图序号"""
        if chapter_num not in self._chapter_counters:
            self._chapter_counters[chapter_num] = {"fig": 0, "tab": 0}

        self._chapter_counters[chapter_num]["fig"] += 1
        return self._chapter_counters[chapter_num]["fig"]

    def _get_next_tab_item_number(self, chapter_num: int) -> int:
        """获取指定章节的下一个表序号"""
        if chapter_num not in self._chapter_counters:
            self._chapter_counters[chapter_num] = {"fig": 0, "tab": 0}

        self._chapter_counters[chapter_num]["tab"] += 1
        return self._chapter_counters[chapter_num]["tab"]

    def _format_caption_static(
        self,
        para,
        caption_type: str,
        number: str,
        title_text: str,
        chinese_font: str = "宋体",
        english_font: str = "Times New Roman",
        caption_size: float = 10.5,
    ) -> None:
        """
        使用静态编号格式化题注，并添加书签支持交叉引用

        生成格式：图/表 章节号-序号 标题
        添加书签：_Ref图1-1 或 _Ref表2-3

        字体规范（学术论文标准）：
        - 中文部分（图/表标签、标题文字）：宋体
        - 英文/数字部分（编号）：Times New Roman

        Args:
            para: 题注段落
            caption_type: '图' 或 '表'
            number: 完整编号（如 "3-1"）
            title_text: 标题文本
            chinese_font: 中文字体名称（默认宋体）
            english_font: 英文字体名称（默认Times New Roman）
            caption_size: 题注字号（默认10.5pt/五号）
        """
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Pt
        except ImportError:
            para.text = f"{caption_type} {number} {title_text}"
            return

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 1.0

        para.clear()

        bookmark_name = f"_Ref{caption_type}{number}"

        run_label = para.add_run(f"{caption_type} ")
        run_label.font.name = english_font
        run_label.font.size = Pt(caption_size)
        rPr = run_label._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), chinese_font)

        self._add_bookmark_start(para, bookmark_name)

        run_number = para.add_run(number)
        run_number.font.name = english_font
        run_number.font.size = Pt(caption_size)

        self._add_bookmark_end(para, bookmark_name)

        if title_text:
            run_space = para.add_run(f" {title_text}")
            run_space.font.name = english_font
            run_space.font.size = Pt(caption_size)
            rPr = run_space._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), chinese_font)

    def _format_caption_with_seq_field(
        self,
        para,
        caption_type: str,
        chapter_num: int,
        item_num: int,
        title_text: str,
        chinese_font: str = "宋体",
        english_font: str = "Times New Roman",
        caption_size: float = 10.5,
    ) -> None:
        """
        使用SEQ域格式化题注，并添加书签支持交叉引用

        生成格式：图/表 章节号-序号 标题
        使用SEQ域：SEQ 图 \\* ARABIC \\s 1
        添加书签：_Ref图1-1 或 _Ref表2-3

        字体规范（学术论文标准）：
        - 中文部分（图/表标签、标题文字）：宋体
        - 英文/数字部分（编号）：Times New Roman

        Args:
            para: 题注段落
            caption_type: '图' 或 '表'
            chapter_num: 章节号
            item_num: 序号
            title_text: 标题文本
            chinese_font: 中文字体名称（默认宋体）
            english_font: 英文字体名称（默认Times New Roman）
            caption_size: 题注字号（默认10.5pt/五号）
        """
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Pt
        except ImportError:
            para.text = f"{caption_type} {chapter_num}-{item_num} {title_text}"
            return

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 1.0

        para.clear()

        full_number = f"{chapter_num}-{item_num}" if chapter_num > 0 else str(item_num)
        bookmark_name = f"_Ref{caption_type}{full_number}"

        run_label = para.add_run(f"{caption_type} ")
        run_label.font.name = english_font
        run_label.font.size = Pt(caption_size)
        rPr = run_label._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from docx.oxml import OxmlElement

            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), chinese_font)

        self._add_bookmark_start(para, bookmark_name)

        if chapter_num > 0:
            run_chap = para.add_run(str(chapter_num))
            run_chap.font.name = english_font
            run_chap.font.size = Pt(caption_size)

            run_sep = para.add_run("-")
            run_sep.font.name = english_font
            run_sep.font.size = Pt(caption_size)

        run_seq = para.add_run()

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = f"SEQ {caption_type} \\* ARABIC \\s 1"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")

        run_placeholder = para.add_run(str(item_num))
        run_placeholder.font.name = english_font
        run_placeholder.font.size = Pt(caption_size)

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")

        run_seq._r.append(fldChar1)
        run_seq._r.append(instrText)
        run_seq._r.append(fldChar2)
        run_seq._r.append(fldChar3)

        self._add_bookmark_end(para, bookmark_name)

        if title_text:
            run_space = para.add_run(f" {title_text}")
            run_space.font.name = english_font
            run_space.font.size = Pt(caption_size)
            rPr = run_space._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                from docx.oxml import OxmlElement

                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), chinese_font)

    def _add_bookmark_start(self, para, bookmark_name: str) -> None:
        """
        在段落当前位置添加书签起始标记（符合 Word OOXML 规范）

        Word 书签结构：
        <w:p>
          <w:bookmarkStart w:id="0" w:name="bookmark_name"/>
          <w:r>内容</w:r>
          <w:bookmarkEnd w:id="0"/>
        </w:p>

        Args:
            para: 段落对象
            bookmark_name: 书签名称
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            bookmark_id = str(abs(hash(bookmark_name)) % 10000)

            bookmark_start = OxmlElement("w:bookmarkStart")
            bookmark_start.set(qn("w:id"), bookmark_id)
            bookmark_start.set(qn("w:name"), bookmark_name)

            para._p.insert(0, bookmark_start)

            para._bookmark_id = bookmark_id
            para._bookmark_name = bookmark_name

        except Exception:
            pass

    def _add_bookmark_end(self, para, bookmark_name: str) -> None:
        """
        在段落末尾添加书签结束标记（符合 Word OOXML 规范）

        Args:
            para: 段落对象
            bookmark_name: 书签名称
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            bookmark_id = getattr(para, "_bookmark_id", str(abs(hash(bookmark_name)) % 10000))

            bookmark_end = OxmlElement("w:bookmarkEnd")
            bookmark_end.set(qn("w:id"), bookmark_id)

            para._p.append(bookmark_end)

        except Exception:
            pass

    def add_fig_pattern(self, pattern: re.Pattern) -> None:
        """
        添加图题注匹配模式

        Args:
            pattern: 正则表达式模式
        """
        self._fig_patterns.append(pattern)

    def add_tab_pattern(self, pattern: re.Pattern) -> None:
        """
        添加表题注匹配模式

        Args:
            pattern: 正则表达式模式
        """
        self._tab_patterns.append(pattern)

    def get_stats(self) -> Dict:
        """
        获取处理统计信息

        Returns:
            统计信息字典
        """
        return {
            "processed_figures": self._processed_figs,
            "processed_tables": self._processed_tabs,
            "skipped_captions": self._skipped_captions,
            "total_processed": self._processed_figs + self._processed_tabs,
            "chapter_counters": dict(self._chapter_counters),
        }
