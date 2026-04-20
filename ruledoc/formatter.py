"""
核心格式化器模块

实现文档加载、处理器链调用、格式应用和保存功能。
"""

import os
from typing import TYPE_CHECKING, Any, List, Optional

from ruledoc.config import get_config
from ruledoc.context import ProcessingContext
from ruledoc.exceptions import ConfigurationError, ProcessingError
from ruledoc.pandoc_converter import PandocConverter, _safe_remove_file, _temp_registry
from ruledoc.processors import CaptionProcessor, HeadingProcessor, ListProcessor, StyleProcessor
from ruledoc.rules.base import FormatRule

if TYPE_CHECKING:
    from docx.document import Document


class Formatter:
    """
    核心格式化器 - 纯流程编排器

    职责:
    - 文档加载 (Markdown → Word 或直接加载 Word)
    - 处理器链调用
    - 规则格式化应用
    - 文档保存

    不负责:
    - 格式定义 (由 FormatRule 提供)
    - 段落检测 (由 FormatRule 提供)
    - 具体转换 (由 PostProcessor 提供)

    Attributes:
        input_path: 输入文件路径
        rule: 格式规则
        output_path: 输出文件路径
        doc: 加载的 Word 文档对象
        converter: Pandoc 转换器
        _temp_files: 临时文件列表 (需要清理)
        _context: 处理上下文
    """

    @property
    def SUPPORTED_INPUT_FORMATS(self) -> set:
        return get_config().file.supported_input_formats

    @property
    def MAX_FILE_SIZE_MB(self) -> int:
        return get_config().file.max_file_size_mb

    @property
    def MAX_FILE_SIZE_BYTES(self) -> int:
        return self.MAX_FILE_SIZE_MB * 1024 * 1024

    def __init__(self, input_path: str, rule: FormatRule, output_path: Optional[str] = None):
        """
        初始化格式化器

        Args:
            input_path: 输入文件路径 (支持 .md, .docx)
            rule: 格式规则实例
            output_path: 输出文件路径 (可选，默认自动生成)

        Raises:
            FileNotFoundError: 输入文件不存在
            ConfigurationError: 不支持的输入格式
        """
        self.input_path = input_path
        self.rule = rule
        self.output_path = output_path or self._generate_output_path()

        self.doc: Any = None
        self.converter = PandocConverter()
        self._temp_files: List[str] = []
        self._context: Optional[ProcessingContext] = None

        self._validate_input()

    def _validate_input(self) -> None:
        """
        验证输入文件

        安全检查:
        - 路径规范化，防止路径遍历攻击
        - 文件大小限制，防止 DoS 攻击
        - 文件格式验证

        Raises:
            ProcessingError: 输入文件不存在或路径不安全
            ConfigurationError: 不支持的输入格式或文件过大
        """
        self.input_path = os.path.abspath(self.input_path)

        if not os.path.exists(self.input_path):
            raise ProcessingError(
                f"输入文件不存在: {self.input_path}", {"input_path": self.input_path}
            )

        real_path = os.path.realpath(self.input_path)
        normalized_path = os.path.normpath(self.input_path)
        if real_path != normalized_path and not os.path.isabs(self.input_path):
            if ".." in self.input_path or not normalized_path.startswith(
                os.path.dirname(real_path)
            ):
                raise ProcessingError(
                    f"输入文件路径包含可疑的路径遍历字符",
                    {"input_path": self.input_path, "real_path": real_path},
                )

        file_size = os.path.getsize(self.input_path)
        if file_size > self.MAX_FILE_SIZE_BYTES:
            raise ConfigurationError(
                f"输入文件过大 ({file_size / 1024 / 1024:.1f}MB)，超过限制 ({self.MAX_FILE_SIZE_MB}MB)",
                {
                    "input_path": self.input_path,
                    "file_size_mb": file_size / 1024 / 1024,
                    "max_size_mb": self.MAX_FILE_SIZE_MB,
                },
            )

        ext = os.path.splitext(self.input_path)[1].lower()
        if ext not in self.SUPPORTED_INPUT_FORMATS:
            raise ConfigurationError(
                f"不支持的输入格式: {ext}",
                {
                    "input_path": self.input_path,
                    "supported_formats": list(self.SUPPORTED_INPUT_FORMATS),
                },
            )

    def _generate_output_path(self) -> str:
        """
        自动生成输出路径

        安全检查:
        - 输出路径规范化
        - 确保输出路径在工作目录内

        Returns:
            输出文件路径
        """
        base = os.path.splitext(self.input_path)[0]
        output_path = f"{base}_formatted.docx"
        output_path = os.path.abspath(output_path)

        cwd = os.path.realpath(os.getcwd())
        real_output = os.path.realpath(output_path)
        if not real_output.startswith(cwd):
            base_name = os.path.basename(output_path)
            output_path = os.path.join(cwd, base_name)

        return output_path

    def _load_document(self) -> Any:
        """
        加载文档

        根据输入文件类型选择加载方式:
        - .md/.txt: 使用 PandocConverter 转换后加载
        - .docx: 直接加载

        Returns:
            python-docx Document 对象

        Raises:
            ProcessingError: 文档加载失败
        """
        ext = os.path.splitext(self.input_path)[1].lower()

        if ext in {".md", ".markdown", ".txt"}:
            return self._load_from_markdown()
        else:
            return self._load_from_word()

    def _load_from_markdown(self) -> Any:
        """
        从 Markdown 文件加载文档

        使用 PandocConverter 将 Markdown 转换为临时 Word 文件，
        然后使用 python-docx 加载。

        Returns:
            python-docx Document 对象

        Raises:
            ProcessingError: 转换或加载失败
        """
        try:
            temp_docx = self.converter.convert_with_temp(self.input_path)
            self._temp_files.append(temp_docx)

            return self._load_word_file(temp_docx)
        except Exception as e:
            raise ProcessingError(
                f"Markdown 文件加载失败: {e}", {"input_path": self.input_path, "error": str(e)}
            )

    def _load_from_word(self) -> Any:
        """
        从 Word 文件加载文档

        Returns:
            python-docx Document 对象

        Raises:
            ProcessingError: 加载失败
        """
        try:
            return self._load_word_file(self.input_path)
        except Exception as e:
            raise ProcessingError(
                f"Word 文件加载失败: {e}", {"input_path": self.input_path, "error": str(e)}
            )

    def _load_word_file(self, file_path: str) -> Any:
        """
        加载 Word 文件

        Args:
            file_path: Word 文件路径

        Returns:
            python-docx Document 对象
        """
        try:
            from docx import Document

            return Document(file_path)
        except ImportError:
            raise ProcessingError(
                "python-docx 未安装，请运行: pip install python-docx",
                {"required_package": "python-docx"},
            )

    def _cleanup_temp_files(self) -> None:
        """
        清理临时文件

        使用安全删除机制，带重试功能。
        """
        for temp_path in self._temp_files:
            if temp_path:
                _safe_remove_file(temp_path)
                _temp_registry.discard(temp_path)
        self._temp_files.clear()

    def process(self) -> None:
        """
        执行格式化流程

        完整流程:
        1. 加载文档
        2. 创建处理上下文
        3. 执行处理器链
        4. 应用规则格式化
        5. 保存文档
        6. 清理临时文件

        Raises:
            ProcessingError: 处理过程中发生错误
        """
        try:
            self.doc = self._load_document()

            self._context = ProcessingContext(doc=self.doc, rule=self.rule)

            self._run_processor_chain(self._context)

            self._apply_rule_formatting(self._context)

            self._save_document()

        except Exception as e:
            if not isinstance(e, (ProcessingError, ConfigurationError)):
                raise ProcessingError(
                    f"格式化处理失败: {e}", {"input_path": self.input_path, "error": str(e)}
                )
            raise
        finally:
            self._cleanup_temp_files()

    def _run_processor_chain(self, ctx: ProcessingContext) -> None:
        """
        执行处理器链

        执行顺序:
        1. 规则前置处理器
        2. 通用处理器 (按 custom_order 或默认顺序: Style → Heading → Caption)
        3. 规则后置处理器

        Args:
            ctx: 处理上下文
        """
        for processor in self.rule.get_pre_processors():
            processor.process(ctx)

        config = self.rule.get_processor_config()

        processors = []
        if config.get("style", True):
            processors.append(("style", StyleProcessor(self.rule.get_style_map())))
        if config.get("heading"):
            if isinstance(config.get("heading"), dict):
                processors.append(("heading", HeadingProcessor(**config.get("heading", {}))))
            else:
                processors.append(("heading", HeadingProcessor()))
        if config.get("list", True):
            processors.append(("list", ListProcessor()))
        if config.get("caption"):
            if isinstance(config.get("caption"), dict):
                processors.append(("caption", CaptionProcessor(**config.get("caption", {}))))
            else:
                processors.append(("caption", CaptionProcessor()))

        order = config.get("custom_order") or ["style", "heading", "list", "caption"]
        for name in order:
            for p_name, processor in processors:
                if p_name == name:
                    processor.process(ctx)

        for processor in self.rule.get_post_processors():
            processor.process(ctx)

    def _apply_rule_formatting(self, ctx: ProcessingContext) -> None:
        """
        应用规则格式化

        Args:
            ctx: 处理上下文
        """
        self._apply_page_settings(ctx)

        self._apply_header_footer(ctx)

        if hasattr(self.rule, "scan_document_structure"):
            self.rule.scan_document_structure(ctx.doc.paragraphs, ctx)

        paragraphs = list(ctx.doc.paragraphs)
        for idx, para in enumerate(paragraphs):
            self._update_section_flags(para, ctx, idx, len(paragraphs))

            para_type = self.rule.detect_paragraph_type(para, ctx)
            self.rule.format_paragraph(para, para_type, ctx)

        if hasattr(self.rule, "format_references_section") and callable(
            getattr(self.rule, "format_references_section")
        ):
            self.rule.format_references_section(ctx.doc.paragraphs, ctx)

        if hasattr(self.rule, "resolve_reference_cross_references") and callable(
            getattr(self.rule, "resolve_reference_cross_references")
        ):
            self.rule.resolve_reference_cross_references(ctx.doc, ctx)  # type: ignore[attr-defined]

        if hasattr(self.rule, "resolve_cross_references") and callable(
            getattr(self.rule, "resolve_cross_references")
        ):
            self.rule.resolve_cross_references(ctx.doc, ctx)  # type: ignore[attr-defined]

        if hasattr(self.rule, "format_tables") and callable(getattr(self.rule, "format_tables")):
            self.rule.format_tables(ctx.doc)  # type: ignore[attr-defined]

        if hasattr(self.rule, "format_formulas") and callable(
            getattr(self.rule, "format_formulas")
        ):
            self.rule.format_formulas(ctx.doc)  # type: ignore[attr-defined]

        if hasattr(self.rule, "format_code_blocks") and callable(
            getattr(self.rule, "format_code_blocks")
        ):
            self.rule.format_code_blocks(ctx.doc)  # type: ignore[attr-defined]

    def _update_section_flags(self, para, ctx: ProcessingContext, _idx: int, _total: int) -> None:
        """
        更新章节标志 (in_abstract, in_references)

        Args:
            para: 当前段落
            ctx: 处理上下文
            _idx: 当前段落索引 (保留供未来扩展)
            _total: 段落总数 (保留供未来扩展)
        """
        text = para.text.strip().lower().replace(" ", "")

        if "摘要" in text or "abstract" in text:
            if len(text) < 20:
                ctx.in_abstract = True
                ctx.in_references = False
                return

        if "关键词" in text or "keywords" in text:
            ctx.in_abstract = False
            return

        if "目录" in text or "contents" in text:
            ctx.in_abstract = False
            return

        # 使用更严格的匹配：必须是独立的标题
        text_no_space = text.replace(" ", "")
        is_ref_heading = text_no_space in ("参考文献", "references") or text in (
            "参考文献",
            "references",
        )
        if is_ref_heading:
            ctx.in_references = True
            ctx.in_abstract = False
            return

        if ctx.in_references:
            if any(
                kw in text
                for kw in ["致谢", "acknowledgement", "附录", "appendix", "结论", "conclusion"]
            ):
                if len(text) < 20:
                    ctx.in_references = False

    def _apply_page_settings(self, ctx: ProcessingContext) -> None:
        """
        应用页面设置

        Args:
            ctx: 处理上下文
        """
        from docx.shared import Cm

        settings = self.rule.get_validated_page_settings()

        for section in ctx.doc.sections:
            if "top_margin" in settings:
                section.top_margin = Cm(settings["top_margin"])
            if "bottom_margin" in settings:
                section.bottom_margin = Cm(settings["bottom_margin"])
            if "left_margin" in settings:
                section.left_margin = Cm(settings["left_margin"])
            if "right_margin" in settings:
                section.right_margin = Cm(settings["right_margin"])
            if "gutter" in settings:
                section.gutter = Cm(settings["gutter"])

    def _apply_header_footer(self, ctx: ProcessingContext) -> None:
        """
        应用页眉页脚设置

        Args:
            ctx: 处理上下文
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt

        settings = self.rule.get_header_footer_settings()

        for section in ctx.doc.sections:
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

            header_text = settings.get("header_text", "")
            if header_text:
                header_para.text = header_text
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for run in header_para.runs:
                    run.font.name = settings.get("header_font", "宋体")
                    run.font.size = Pt(settings.get("header_font_size", 9))
                    rPr = run._element.rPr
                    if rPr is not None:
                        rFonts = rPr.find(qn("w:rFonts"))
                        if rFonts is not None:
                            rFonts.set(qn("w:eastAsia"), settings.get("header_font", "宋体"))

            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if settings.get("footer_type") == "page_number":
                self._add_page_number(footer_para, settings)

    def _add_page_number(self, paragraph, settings: dict) -> None:
        """
        添加页码字段

        Args:
            paragraph: 段落对象
            settings: 页眉页脚设置
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        run = paragraph.add_run()

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        run.font.name = "Times New Roman"
        run.font.size = Pt(settings.get("header_font_size", 9))

    def _save_document(self) -> None:
        """
        保存文档

        Raises:
            ProcessingError: 保存失败
        """
        if self.doc is None:
            raise ProcessingError("文档未加载，无法保存", {"output_path": self.output_path})

        try:
            output_dir = os.path.dirname(self.output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)

            self.doc.save(self.output_path)
        except Exception as e:
            raise ProcessingError(
                f"文档保存失败: {e}", {"output_path": self.output_path, "error": str(e)}
            )

    def get_warnings(self) -> List[str]:
        """
        获取处理警告

        Returns:
            警告消息列表
        """
        if self._context is not None:
            return self._context.warnings
        return []
