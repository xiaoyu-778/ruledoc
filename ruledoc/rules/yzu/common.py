"""扬州大学 (YZU) 规则共享模块

包含常量定义和共享方法。
"""

import re
from enum import Enum
from typing import TYPE_CHECKING, Dict, List, Optional

from ruledoc.config import get_config

if TYPE_CHECKING:
    pass


class ThesisType(Enum):
    """论文类型"""

    THESIS = "毕业论文"
    DESIGN_REPORT = "毕业设计报告"


class YZUConstants:
    """扬州大学论文格式常量"""

    @property
    def PAGE_SETTINGS(self) -> Dict[str, float]:
        config = get_config()
        return {
            "top_margin": config.page.top_margin,
            "bottom_margin": config.page.bottom_margin,
            "left_margin": config.page.left_margin,
            "right_margin": config.page.right_margin,
            "gutter": config.page.gutter,
            "header_distance": config.page.header_distance,
            "footer_distance": config.page.footer_distance,
        }

    @property
    def FONT_SETTINGS(self) -> Dict[str, str]:
        config = get_config()
        return {
            "chinese_font": config.fonts.chinese,
            "english_font": config.fonts.english,
            "heading_font": config.fonts.heading,
            "abstract_font_thesis": config.fonts.abstract_thesis,
            "abstract_font_design": config.fonts.abstract_design,
            "signature_font": config.fonts.signature,
        }

    @property
    def FONT_SIZES(self) -> Dict[str, float]:
        config = get_config()
        return {
            "title": config.font_sizes.title,
            "heading_1": config.font_sizes.heading_1,
            "heading_2": config.font_sizes.heading_2,
            "heading_3": config.font_sizes.heading_3,
            "heading_4": config.font_sizes.heading_4,
            "body": config.font_sizes.body,
            "caption": config.font_sizes.caption,
            "header_footer": config.font_sizes.header_footer,
        }

    @property
    def LINE_SPACING(self) -> float:
        return get_config().paragraph.line_spacing

    @property
    def TAB_CENTER_POS_CM(self) -> float:
        return get_config().tab_stops.center

    @property
    def TAB_RIGHT_POS_CM(self) -> float:
        return get_config().tab_stops.right

    @property
    def SPECIAL_TITLES(self) -> set:
        return get_config().special_titles.all

    @property
    def SIGNATURE_KEYWORDS(self) -> List[str]:
        return get_config().signature_keywords.items


YZU_CONSTANTS = YZUConstants()


class YZUMixin:
    """扬州大学规则共享方法 Mixin"""

    @property
    def PAGE_SETTINGS(self) -> Dict[str, float]:
        return YZU_CONSTANTS.PAGE_SETTINGS

    @property
    def FONT_SETTINGS(self) -> Dict[str, str]:
        return YZU_CONSTANTS.FONT_SETTINGS

    @property
    def FONT_SIZES(self) -> Dict[str, float]:
        return YZU_CONSTANTS.FONT_SIZES

    @property
    def LINE_SPACING(self) -> float:
        return YZU_CONSTANTS.LINE_SPACING

    @property
    def TAB_CENTER_POS_CM(self) -> float:
        return YZU_CONSTANTS.TAB_CENTER_POS_CM

    @property
    def TAB_RIGHT_POS_CM(self) -> float:
        return YZU_CONSTANTS.TAB_RIGHT_POS_CM

    @property
    def SPECIAL_TITLES(self) -> set:
        return YZU_CONSTANTS.SPECIAL_TITLES

    @property
    def SIGNATURE_KEYWORDS(self) -> List[str]:
        return YZU_CONSTANTS.SIGNATURE_KEYWORDS

    def _set_east_asian_font(self, run, font_name: str) -> None:
        """安全设置东亚字体（符合 Word OOXML 规范）

        Word 字体设置需要同时指定：
        - w:ascii: ASCII 字符（英文、数字）
        - w:hAnsi: 高位 ANSI 字符
        - w:eastAsia: 东亚字符（中文、日文、韩文）

        Args:
            run: Run 对象
            font_name: 东亚字体名称
        """
        try:
            from docx.oxml.ns import qn

            if run._element is None:
                return
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                rFonts.set(qn("w:eastAsia"), font_name)
        except Exception as e:
            import logging

            logging.debug(f"YZUMixin: 设置东亚字体失败: {e}")

    def _add_text_run(
        self, para, text: str, use_chinese_font: bool = False, is_superscript: bool = False
    ) -> None:
        """添加带有正确字体设置的文本 run

        Args:
            para: 段落对象
            text: 文本内容
            use_chinese_font: 是否使用中文字体
            is_superscript: 是否为上标
        """
        try:
            from docx.shared import Pt, RGBColor
        except ImportError:
            para.add_run(text)
            return

        run = para.add_run(text)
        run.font.name = self.FONT_SETTINGS["english_font"]
        run.font.size = Pt(self.FONT_SIZES["body"])
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.superscript = is_superscript

        if use_chinese_font:
            self._set_east_asian_font(run, self.FONT_SETTINGS["chinese_font"])

    def _build_ref_field(
        self, para, bookmark_name: str, display_text: str, is_superscript: bool = False
    ) -> None:
        """构建 REF 域并添加到段落

        生成 Word REF 域结构：
        - fldChar begin
        - instrText "REF bookmark_name \\h"
        - fldChar separate
        - placeholder text
        - fldChar end

        Args:
            para: 段落对象
            bookmark_name: 书签名称
            display_text: 显示文本
            is_superscript: 是否为上标
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Pt, RGBColor
        except ImportError:
            run = para.add_run(display_text)
            run.font.superscript = is_superscript
            return

        run_begin = para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run_begin._r.append(fldChar1)

        run_instr = para.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = f"REF {bookmark_name} \\h"
        run_instr._r.append(instrText)

        run_sep = para.add_run()
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        run_sep._r.append(fldChar2)

        run_placeholder = para.add_run(display_text)
        run_placeholder.font.name = self.FONT_SETTINGS["english_font"]
        run_placeholder.font.size = Pt(self.FONT_SIZES["body"])
        run_placeholder.font.color.rgb = RGBColor(0, 0, 0)
        run_placeholder.font.superscript = is_superscript
        # 设置东亚字体（中文部分使用宋体）
        if self._has_chinese(display_text):
            self._set_east_asian_font(run_placeholder, self.FONT_SETTINGS["chinese_font"])

        run_end = para.add_run()
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run_end._r.append(fldChar3)

    def _has_chinese(self, text: str) -> bool:
        """检查文本是否包含中文字符"""
        return any("\u4e00" <= char <= "\u9fff" for char in text)

    def _parse_citation_numbers(self, citation_raw: str) -> List[int]:
        """解析引用编号字符串为独立的数字列表

        支持格式：
        - "4" → [4]
        - "1-3" → [1, 2, 3]
        - "1,2,3" → [1, 2, 3]
        - "1,3-5" → [1, 3, 4, 5]

        Args:
            citation_raw: 原始引用编号字符串

        Returns:
            排序后的唯一数字列表
        """
        numbers = set()
        parts = citation_raw.split(",")

        for part in parts:
            part = part.strip()
            if "-" in part:
                range_parts = part.split("-")
                try:
                    start = int(range_parts[0].strip())
                    end = int(range_parts[1].strip())
                    for n in range(start, end + 1):
                        numbers.add(n)
                except ValueError:
                    pass
            else:
                try:
                    numbers.add(int(part))
                except ValueError:
                    pass

        return sorted(numbers)

    def _find_replacement_number(self, old_ref: str, nums_list: list) -> Optional[str]:
        """查找对应的实际编号"""
        if old_ref in nums_list:
            return old_ref

        try:
            parts = old_ref.replace("–", "-").replace("—", "-").replace("‐", "-").split("-")
            if len(parts) == 2:
                ref_seq = int(parts[1])

                for actual_num in nums_list:
                    act_parts = actual_num.split("-")
                    if len(act_parts) == 2:
                        if int(act_parts[1]) == ref_seq:
                            return actual_num
        except (ValueError, IndexError):
            pass

        try:
            seq = int(old_ref.split("-")[-1])
            if 0 < seq <= len(nums_list):
                return nums_list[seq - 1]
        except (ValueError, IndexError):
            pass

        return None

    def _collect_caption_numbers(self, doc, label: str) -> list:
        """收集文档中所有指定类型的题注编号"""
        pattern = re.compile(rf"^{re.escape(label)}\s+(\d+[-‐–—]\d+)", re.IGNORECASE)
        numbers = []
        for para in doc.paragraphs:
            m = pattern.match(para.text.strip())
            if m:
                numbers.append(m.group(1))
        return numbers

    def _collect_formula_numbers(self, doc) -> list:
        """收集文档中所有公式的章节化编号"""
        numbers = []
        try:
            from docx.oxml.ns import qn
        except ImportError:
            return numbers

        formula_pattern = re.compile(r"^\s*\(?\d+[-‐–—]\d+\)?")

        for para in doc.paragraphs:
            has_math = False
            try:
                xml = para._p.xml
                if "m:oMath" in xml or "oMath" in xml:
                    has_math = True
            except Exception:
                pass
            if not has_math:
                continue
            if self._is_inline_formula(para):
                continue

            m = formula_pattern.match(para.text.strip())
            if m:
                num_text = m.group(0).strip().strip("()")
                normalized = num_text.replace("–", "-").replace("—", "-").replace("‐", "-")
                numbers.append(normalized)
        return numbers

    def _is_inline_formula(self, para) -> bool:
        """判断是否为行内公式"""
        try:
            xml = para._p.xml
        except Exception:
            return False

        text = para.text.strip()

        number_pattern = re.compile(r"^\s*\(?\d+[-‐–—]?\d*\)?\s*$")
        if number_pattern.match(text):
            return False

        all_runs_text = "".join(run.text for run in para.runs).strip()

        non_math_text = re.sub(r"[()\d\s\-‐–—．.]", "", all_runs_text)

        if len(non_math_text) > 3:
            return True

        run_texts = [run.text.strip() for run in para.runs if run.text.strip()]
        non_empty_runs = [t for t in run_texts if t]

        if len(non_empty_runs) > 2:
            return True

        non_empty_texts = [
            t for t in non_empty_runs if not re.match(r"^[\s(]\d+[-‐–—]?\d*[\)]?\s*$", t)
        ]
        for t in non_empty_texts:
            if len(t) > 2 and not re.match(r"^[()\d\s\-‐–—．.]+$", t):
                return True

        return False

    def _run_contains_math(self, run) -> bool:
        """检测 Run 对象是否包含 OMML 数学公式"""
        try:
            xml = run._r.xml if hasattr(run, "_r") and run._r is not None else ""
            return "m:oMath" in xml or "oMath" in xml
        except Exception:
            return False
