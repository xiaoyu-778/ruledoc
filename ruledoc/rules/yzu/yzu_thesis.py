"""扬州大学毕业论文格式规则

基于《扬州大学本科生毕业设计（论文）格式及要求》实现。
适用于毕业论文类型文档。
"""

import re
from typing import TYPE_CHECKING, Dict, List, Tuple

from ruledoc.rules.base import FormatRule, register_rule

from .common import ThesisType, YZUMixin

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

    from ruledoc.context import ProcessingContext


@register_rule
class YZUThesisRule(YZUMixin, FormatRule):
    """扬州大学毕业论文格式规则

    符合扬州大学本科生毕业设计（论文）格式及要求。
    适用于毕业论文类型文档。

    Attributes:
        thesis_type: 论文类型 (毕业论文/毕业设计报告)

    命名说明:
        - YZU: Yangzhou University (扬州大学)
        - Thesis: 学位论文
        - Rule: 规则类后缀
    """

    def __init__(self, thesis_type: ThesisType = ThesisType.THESIS):
        """初始化 YZU 规则

        Args:
            thesis_type: 论文类型 (毕业论文/毕业设计报告)
        """
        self.thesis_type = thesis_type
        self._title_detected = False

    @property
    def name(self) -> str:
        return "yzu_thesis"

    @property
    def description(self) -> str:
        return f"扬州大学{self.thesis_type.value}格式"

    def get_page_settings(self) -> Dict[str, float]:
        return self.PAGE_SETTINGS.copy()

    def get_font_settings(self) -> Dict[str, str]:
        settings = self.FONT_SETTINGS.copy()
        settings["abstract_font"] = self.get_abstract_font()
        return settings

    def get_heading_format(self, level: int) -> Tuple[str, str, float]:
        """获取标题格式

        关键差异:
        - 毕业论文一级标题居中
        - 毕业设计报告一级标题靠左
        """
        if level == 1:
            alignment = "center" if self.thesis_type == ThesisType.THESIS else "left"
        else:
            alignment = "left"

        formats = {
            1: (self.FONT_SETTINGS["heading_font"], alignment, self.FONT_SIZES["heading_1"]),
            2: (self.FONT_SETTINGS["heading_font"], "left", self.FONT_SIZES["heading_2"]),
            3: (self.FONT_SETTINGS["heading_font"], "left", self.FONT_SIZES["heading_3"]),
            4: (self.FONT_SETTINGS["heading_font"], "left", self.FONT_SIZES["heading_4"]),
        }
        return formats.get(level, (self.FONT_SETTINGS["heading_font"], "left", 12))

    def get_abstract_font(self) -> str:
        """获取摘要内容字体"""
        if self.thesis_type == ThesisType.THESIS:
            return self.FONT_SETTINGS["abstract_font_thesis"]
        return self.FONT_SETTINGS["abstract_font_design"]

    def get_processor_config(self) -> Dict:
        return {
            "style": True,
            "heading": True,
            "caption": True,
            "custom_order": None,
        }

    def get_header_footer_settings(self) -> Dict:
        return {
            "header_text": "扬州大学本科生毕业论文",
            "header_font": self.FONT_SETTINGS["chinese_font"],
            "header_font_size": self.FONT_SIZES["header_footer"],
            "footer_type": "page_number",
        }

    def detect_paragraph_type(self, para: "Paragraph", ctx: "ProcessingContext") -> str:
        """检测段落类型

        类型包括:
        - title: 论文标题（第一个 Heading 1，独立于一级标题）
        - heading: 标题
        - caption: 题注
        - abstract_heading: 摘要标题
        - abstract: 摘要内容
        - references_heading: 参考文献标题
        - references: 参考文献条目
        - signature: 署名
        - special_title: 特殊标题
        - body: 正文
        """
        text = para.text.strip()

        if not text:
            return "empty"

        if self._is_title_paragraph(para, ctx):
            return "title"

        if self._is_special_title(text):
            return "special_title"

        if self._is_heading(para):
            return "heading"

        if self._is_signature(text, ctx):
            return "signature"

        if self._is_caption(text):
            return "caption"

        if self._is_abstract_content(para, ctx):
            return "abstract"

        if self._is_references_content(para, ctx):
            return "references"

        return "body"

    def _is_heading(self, para: "Paragraph") -> bool:
        """检测是否为标题（排除已识别为论文标题的 Heading 1）"""
        if not para.style:
            return False
        style_name = para.style.name.lower() if para.style.name else ""
        if not style_name.startswith("heading"):
            return False

        # 排除特殊标题（通过文本内容判断）
        text = para.text.strip().lower().replace(" ", "")
        if self._is_special_title(text) and "引言" not in text and "introduction" not in text:
            # 引言虽然在 SpecialTitles 中，但它属于 Heading 层面的一种（通常是第一章）
            # 这里如果不包含引言，就排除掉
            return False

        return True

    def _is_title_paragraph(self, para: "Paragraph", ctx: "ProcessingContext") -> bool:
        """检测是否为论文标题
        """
        # 如果已经标记过标题段落，直接对比 identity
        if hasattr(ctx, "_title_para_id"):
            return id(para) == ctx._title_para_id

        if not para.style:
            return False

        style_name = para.style.name.lower() if para.style.name else ""

        # 1. 优先使用 Title 样式
        if style_name == "title" or style_name == "标题":
            ctx._title_para_id = id(para)
            return True

        # 2. 回退：文档中第一个 H1，且不是特定关键词标题（摘要、参考文献等）
        if style_name == "heading 1" or style_name == "heading1":
            text = para.text.strip().lower().replace(" ", "")
            if not self._is_special_title(text):
                ctx._title_para_id = id(para)
                return True

        return False

    def _is_special_title(self, text: str) -> bool:
        """检测是否为特殊标题

        使用精确匹配或开头匹配，避免将包含关键词的正文误判为标题。
        例如："确认参考文献引用格式是否正确" 不应被识别为参考文献标题。
        """
        text_stripped = text.strip()
        text_lower = text_stripped.lower().replace(" ", "")
        text_len = len(text_stripped)

        # 只检查短文本（标题通常较短）
        if text_len > 20:
            return False

        for title in self.SPECIAL_TITLES:
            # 精确匹配：文本完全等于关键词
            if text_lower == title:
                return True

            # 开头匹配：文本以关键词开头，后跟标点或空格
            if text_lower.startswith(title):
                remaining = text_lower[len(title) :]
                if not remaining or remaining[0] in r":\s,;，；":
                    return True

        return False

    def _is_signature(self, text: str, _ctx: "ProcessingContext") -> bool:
        """检测是否为署名"""
        count = sum(1 for k in self.SIGNATURE_KEYWORDS if k in text)
        if count >= 2:
            return True

        if re.match(r"^(学生|教师|专业|班级|学号)[:：]", text):
            return True

        return False

    def _is_caption(self, text: str) -> bool:
        """检测是否为题注

        严格的题注检测规则：
        1. 必须以"图/表/Figure/Table"开头，后跟数字编号
        2. 文本长度限制：题注通常较短（不超过100字符）
        3. 必须包含标题内容（不能只有"图 1-1"）
        4. 不能是正文中的引用（如"图 1-1 显示..."）

        Args:
            text: 段落文本

        Returns:
            是否为题注
        """
        # 长度检查：题注通常较短，正文引用通常较长
        if len(text) > 100:
            return False

        # 匹配"图 1-1 标题"或"Figure 1-1 Title"格式
        # 要求：前缀 + 编号 + 空格 + 标题内容
        caption_pattern = re.compile(
            r"^(图|表|Figure|Table|Fig\.?|Tab\.?)\s*\d+[-‐–—]?\d*\s+[^\s].+$", re.IGNORECASE
        )

        if not caption_pattern.match(text):
            return False

        # 排除正文引用模式："图 X-Y 是/显示/表明/如图..."
        # 如果"图/表"后面紧跟着是动词或常用正文连接词，则不是题注
        reference_patterns = [
            re.compile(
                r"^(图|表)\s*\d+[-‐–—]?\d*\s+(是|显示|表明|展示|说明|如图|见|为|可以|在)",
                re.IGNORECASE,
            ),
            re.compile(
                r"^(Figure|Table|Fig\.?|Tab\.?)\s*\d+[-‐–—]?\d*\s+(is|shows?|illustrates?|demonstrates?|as)",
                re.IGNORECASE,
            ),
        ]

        for pattern in reference_patterns:
            if pattern.match(text):
                return False

        return True

    def _is_abstract_content(self, _para: "Paragraph", ctx: "ProcessingContext") -> bool:
        """检测是否为摘要内容"""
        return ctx.in_abstract

    def _is_references_content(self, _para: "Paragraph", ctx: "ProcessingContext") -> bool:
        """检测是否为参考文献内容"""
        if ctx.in_references:
            return True
        return False

    def scan_document_structure(
        self, paragraphs: List["Paragraph"], ctx: "ProcessingContext"
    ) -> None:
        """扫描文档结构（已弃用，保留用于向后兼容）

        注意: in_abstract 和 in_references 标志现在由 Formatter._update_section_flags 设置

        Args:
            paragraphs: 段落列表
            ctx: 处理上下文
        """
        _ = paragraphs
        _ = ctx
        pass

    def format_paragraph(self, para: "Paragraph", para_type: str, ctx: "ProcessingContext") -> None:
        """格式化段落

        Args:
            para: 段落对象
            para_type: 段落类型
            ctx: 处理上下文
        """
        if para_type == "empty":
            return

        if para_type == "title":
            self._format_title(para)
        elif para_type == "heading":
            self._format_heading(para, ctx)
        elif para_type == "special_title":
            text = para.text.strip()
            text_lower = text.lower().replace(" ", "")
            if "关键词" in text_lower or "keywords" in text_lower:
                self._format_keywords(para)
            else:
                self._format_special_title(para)
        elif para_type == "signature":
            self._format_signature(para)
        elif para_type == "caption":
            pass
        elif para_type == "abstract":
            self._format_abstract_content(para)
        elif para_type == "references":
            self._format_reference_item(para)
        else:
            self._format_body(para)

    def _format_title(self, para: "Paragraph") -> None:
        """格式化论文标题（level=0，独立于一级标题）

        参考 old 脚本: level=0 使用黑体18pt居中，不应用多级列表编号

        论文标题格式要求：
        - 字体：黑体
        - 字号：小二号（18pt）
        - 位置：居中
        - 不编号
        - 大纲级别设为正文文本（非标题）
        """
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt, RGBColor
        except ImportError:
            return

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        from docx.oxml.ns import qn

        pPr = para._p.pPr
        if pPr is not None:
            spacing = pPr.spacing
            if spacing is not None:
                spacing.set(qn("w:beforeLines"), "100")
                spacing.set(qn("w:afterLines"), "100")

        self._set_outline_level_to_body(para)

        for run in para.runs:
            run.font.name = self.FONT_SETTINGS["english_font"]
            run.font.size = Pt(self.FONT_SIZES["title"])
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_east_asian_font(run, self.FONT_SETTINGS["heading_font"])

    def _format_heading(self, para: "Paragraph", _ctx: "ProcessingContext") -> None:
        """格式化标题（段前段后1行=字号×行距，一级标题段前分页）"""
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Cm, Pt, RGBColor
        except ImportError:
            return

        style_name = para.style.name.lower() if para.style and para.style.name else ""
        level = 0
        for i in range(1, 5):
            if f"heading {i}" in style_name:
                level = i
                break

        if level == 0:
            return

        font_name, alignment, font_size = self.get_heading_format(level)

        if alignment == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        para.paragraph_format.line_spacing = 1.0

        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        from docx.oxml.ns import qn

        pPr = para._p.pPr
        if pPr is not None:
            spacing = pPr.spacing
            if spacing is not None:
                spacing.set(qn("w:beforeLines"), "100")
                spacing.set(qn("w:afterLines"), "100")

        if level == 1:
            para.paragraph_format.page_break_before = True

        for run in para.runs:
            run.font.name = self.FONT_SETTINGS["english_font"]
            run.font.size = Pt(font_size)
            run.font.bold = False
            run.font.italic = False
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_east_asian_font(run, font_name)

    def _format_special_title(self, para: "Paragraph") -> None:
        """格式化特殊标题（摘要、目录、致谢、参考文献、附录等）

        格式要求：
        - 对齐方式：居中
        - 字体：黑体
        - 字号：
          - 摘要：小四（12pt）
          - 目录、致谢、参考文献、附录等：小三（15pt）
        - 不加编号
        - 摘要不设置段前分页，其他特殊标题设置段前分页
        - 摘要的大纲级别设为正文文本（非标题）
        """
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Pt, RGBColor
        except ImportError:
            return

        text = para.text.strip().lower().replace(" ", "")

        is_abstract = "摘要" in text or "abstract" in text
        is_intro = "引言" in text or "introduction" in text

        # 摘要字号小四 (12pt)，其他特殊标题小三 (15pt)
        font_size = self.FONT_SIZES["body"] if is_abstract else self.FONT_SIZES["heading_1"]

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 1.0

        if is_abstract:
            para.paragraph_format.page_break_before = False
            # 摘要设置为正文大纲级别
            self._set_outline_level_to_body(para)
        else:
            # 引言默认也要页前分页
            para.paragraph_format.page_break_before = True

        # 间距：1 行
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        pPr = para._p.pPr
        if pPr is not None:
            spacing = pPr.spacing
            if spacing is not None:
                spacing.set(qn("w:beforeLines"), "100")
                spacing.set(qn("w:afterLines"), "100")

        for run in para.runs:
            run.font.name = self.FONT_SETTINGS["english_font"]
            run.font.size = Pt(font_size)
            # 摘要不加粗，其他特殊标题加粗
            run.font.bold = not is_abstract
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_east_asian_font(run, self.FONT_SETTINGS["heading_font"])

    def _set_outline_level_to_body(self, para: "Paragraph") -> None:
        """将段落的大纲级别设置为正文文本（非标题）

        通过移除 pStyle 中的 Heading 样式，使段落不被视为标题。
        这对于摘要等特殊标题很重要，因为它们不应该出现在目录中。

        Args:
            para: 段落对象
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ImportError:
            return

        pPr = para._p.get_or_add_pPr()

        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            pPr.remove(pStyle)

        outlineLvl = pPr.find(qn("w:outlineLvl"))
        if outlineLvl is not None:
            pPr.remove(outlineLvl)

    def _format_keywords(self, para: "Paragraph") -> None:
        """格式化关键词段落

        关键词格式要求：
        - 标签"关键词："：黑体小四号
        - 关键词内容：中文宋体 + 英文Times New Roman 小四号（与正文相同）
        - 对齐方式：左对齐

        段落结构示例：
        "关键词：格式化；测试文档；RuleDoc；论文"
        或
        "Keywords: formatting; test document; RuleDoc; thesis"
        """
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt, RGBColor
        except ImportError:
            return

        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        para.paragraph_format.line_spacing = 1.0

        text = para.text
        if not text:
            return

        label_match = None
        for pattern in [r"^(关键词\s*[:：])", r"^(Keywords?\s*[:])"]:
            match = __import__("re").match(pattern, text, __import__("re").IGNORECASE)
            if match:
                label_match = match
                break

        if label_match:
            label_text = label_match.group(1)
            content_text = text[label_match.end() :]

            para.clear()

            label_run = para.add_run(label_text)
            label_run.font.name = self.FONT_SETTINGS["english_font"]
            label_run.font.size = Pt(self.FONT_SIZES["body"])
            label_run.font.bold = True
            label_run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_east_asian_font(label_run, self.FONT_SETTINGS["heading_font"])

            if content_text:
                content_run = para.add_run(content_text)
                content_run.font.name = self.FONT_SETTINGS["english_font"]
                content_run.font.size = Pt(self.FONT_SIZES["body"])
                content_run.font.bold = False
                content_run.font.color.rgb = RGBColor(0, 0, 0)

                if any("\u4e00" <= char <= "\u9fff" for char in content_text):
                    self._set_east_asian_font(content_run, self.FONT_SETTINGS["chinese_font"])
        else:
            for run in para.runs:
                run.font.name = self.FONT_SETTINGS["english_font"]
                run.font.size = Pt(self.FONT_SIZES["body"])
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                self._set_east_asian_font(run, self.FONT_SETTINGS["heading_font"])

    def _format_signature(self, para: "Paragraph") -> None:
        """格式化署名（个人信息）

        字体要求：
        - 中文：楷体小四（12pt）
        - 英文和数字：Times New Roman 小四（12pt）
        """
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt, RGBColor
        except ImportError:
            return

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = self.LINE_SPACING

        for run in para.runs:
            run.font.name = self.FONT_SETTINGS["english_font"]
            run.font.size = Pt(self.FONT_SIZES["body"])
            run.font.color.rgb = RGBColor(0, 0, 0)

            run_text = run.text
            if run_text and any("\u4e00" <= char <= "\u9fff" for char in run_text):
                self._set_east_asian_font(run, self.FONT_SETTINGS["signature_font"])

    def _format_abstract_content(self, para: "Paragraph") -> None:
        """格式化摘要内容"""
        try:
            from docx.shared import Cm, Pt, RGBColor
        except ImportError:
            return

        para.paragraph_format.line_spacing = self.LINE_SPACING
        para.paragraph_format.first_line_indent = Cm(0.74)

        abstract_font = self.get_abstract_font()

        for run in para.runs:
            run.font.name = self.FONT_SETTINGS["english_font"]
            run.font.size = Pt(self.FONT_SIZES["body"])
            run.font.color.rgb = RGBColor(0, 0, 0)
            text = run.text
            if text and any("\u4e00" <= char <= "\u9fff" for char in text):
                self._set_east_asian_font(run, abstract_font)

    def _format_reference_item(self, para: "Paragraph") -> None:
        """格式化参考文献条目"""
        try:
            from docx.shared import Cm, Pt
        except ImportError:
            return

        para.paragraph_format.line_spacing = self.LINE_SPACING
        para.paragraph_format.first_line_indent = Cm(-0.74)
        para.paragraph_format.left_indent = Cm(0.74)

        self._apply_mixed_fonts(para)

    NUMBERING_PATTERNS = [
        re.compile(r"^(\d+[\.\、．])\s*"),
        re.compile(r"^([（\(]\d+[）\)])\s*"),
        re.compile(r"^([（\(][一二三四五六七八九十]+[）\)])\s*"),
        re.compile(r"^([一二三四五六七八九十]+[\.\、．])\s*"),
        re.compile(r"^([①②③④⑤⑥⑦⑧⑨⑩])\s*"),
    ]

    def _format_body(self, para: "Paragraph") -> None:
        """格式化正文（段前段后0磅）

        支持正文中手写编号的识别和格式化：
        - 阿拉伯数字编号：1. 2. 3. 或 1、2、3、
        - 带括号编号：（1）（2）（3）或 (1)(2)(3)
        - 中文数字编号：一、二、三、 或 （一）（二）（三）
        - 圆圈数字：①②③

        编号使用 Times New Roman 小四号字体，正文中文使用宋体小四号。
        带编号的段落不设置首行缩进，避免编号错位。
        """
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Pt, RGBColor
        except ImportError:
            return

        # 明确设置左对齐，防止继承其他对齐方式（如居中）
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        pf = para.paragraph_format
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        self._process_reference_citations(para)

        text = para.text
        numbering_match = None
        for pattern in self.NUMBERING_PATTERNS:
            match = pattern.match(text)
            if match:
                numbering_match = match
                break

        if numbering_match:
            self._format_body_with_numbering(para, numbering_match)
        else:
            self._set_first_line_indent_chars(para, 2)

            pf.left_indent = Pt(0)
            pf.right_indent = Pt(0)

            for run in para.runs:
                run.font.name = self.FONT_SETTINGS["english_font"]
                run.font.size = Pt(self.FONT_SIZES["body"])
                run.font.bold = False
                run.font.italic = False
                run.font.underline = False
                run.font.color.rgb = RGBColor(0, 0, 0)

                run_text = run.text
                if run_text and any("\u4e00" <= char <= "\u9fff" for char in run_text):
                    self._set_east_asian_font(run, self.FONT_SETTINGS["chinese_font"])

    def _format_body_with_numbering(self, para: "Paragraph", numbering_match) -> None:
        """格式化带有手写编号的正文段落

        编号部分使用 Times New Roman 小四号字体，
        正文部分中文使用宋体，英文/数字使用 Times New Roman。

        带编号段落不设置首行缩进，避免编号错位。

        Args:
            para: 段落对象
            numbering_match: 编号正则匹配对象
        """
        try:
            from docx.shared import Pt, RGBColor
        except ImportError:
            return

        numbering_text = numbering_match.group(1)
        numbering_end = numbering_match.end()
        full_text = para.text
        body_text = full_text[numbering_end:]

        para.clear()

        num_run = para.add_run(numbering_text)
        num_run.font.name = self.FONT_SETTINGS["english_font"]
        num_run.font.size = Pt(self.FONT_SIZES["body"])
        num_run.font.bold = False
        num_run.font.color.rgb = RGBColor(0, 0, 0)

        if any("\u4e00" <= char <= "\u9fff" for char in numbering_text):
            self._set_east_asian_font(num_run, self.FONT_SETTINGS["chinese_font"])

        if body_text:
            self._add_body_text_with_mixed_fonts(para, body_text)

    def _add_body_text_with_mixed_fonts(self, para, text: str) -> None:
        """添加正文文本，中英文使用不同字体

        中文使用宋体，英文和数字使用 Times New Roman。
        字号统一为小四号（12pt）。

        Args:
            para: 段落对象
            text: 文本内容
        """
        try:
            from docx.shared import Pt, RGBColor
        except ImportError:
            para.add_run(text)
            return

        current_text = ""
        is_current_chinese = None

        for char in text:
            is_chinese = "\u4e00" <= char <= "\u9fff"

            if is_current_chinese is None:
                is_current_chinese = is_chinese
                current_text = char
            elif is_current_chinese == is_chinese:
                current_text += char
            else:
                if current_text:
                    run = para.add_run(current_text)
                    run.font.size = Pt(self.FONT_SIZES["body"])
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    if is_current_chinese:
                        run.font.name = self.FONT_SETTINGS["chinese_font"]
                        self._set_east_asian_font(run, self.FONT_SETTINGS["chinese_font"])
                    else:
                        run.font.name = self.FONT_SETTINGS["english_font"]
                is_current_chinese = is_chinese
                current_text = char

        if current_text:
            run = para.add_run(current_text)
            run.font.size = Pt(self.FONT_SIZES["body"])
            run.font.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0)
            if is_current_chinese:
                run.font.name = self.FONT_SETTINGS["chinese_font"]
                self._set_east_asian_font(run, self.FONT_SETTINGS["chinese_font"])
            else:
                run.font.name = self.FONT_SETTINGS["english_font"]

    def _set_first_line_indent_chars(self, para, chars: int) -> None:
        """设置首行缩进（字符单位）

        Word OOXML 中使用 w:firstLineChars 属性，单位为 1/100 字符。
        例如：2 字符 = 200

        Args:
            para: 段落对象
            chars: 缩进字符数
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ImportError:
            return

        pPr = para._p.get_or_add_pPr()

        # 查找或创建 w:ind 元素
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)

        # 设置首行缩进字符数（单位：1/100 字符）
        ind.set(qn("w:firstLineChars"), str(chars * 100))
        # 同时设置 firstLine（兼容性，以 twips 为单位，约等于字符宽度）
        # 1 字符约等于字号（pt），转换为 twips: pt * 20
        # 这里假设正文字号为12pt (小四) -> 240 twips / 字符
        ind.set(qn("w:firstLine"), str(chars * 240))

    def _process_reference_citations(self, para: "Paragraph") -> None:
        """处理正文中的参考文献引用，转换为上标格式

        支持格式:
        - [1], [2], [3] 单个引用
        - [1-3] 连续引用
        - [1,2,3] 多个引用
        - ^[1]^ Markdown格式
        """
        try:
            from docx.shared import RGBColor
        except ImportError:
            return

        text = para.text
        if not text:
            return

        citation_pattern = re.compile(r"\^?\[(\d+(?:[-,]\d+)*)\]\^?")

        matches = list(citation_pattern.finditer(text))
        if not matches:
            return

        para.clear()

        last_end = 0
        for match in matches:
            if match.start() > last_end:
                normal_text = text[last_end : match.start()]
                run = para.add_run(normal_text)
                run.font.name = self.FONT_SETTINGS["english_font"]
                run.font.color.rgb = RGBColor(0, 0, 0)
                if any("\u4e00" <= char <= "\u9fff" for char in normal_text):
                    self._set_east_asian_font(run, self.FONT_SETTINGS["chinese_font"])

            citation_text = f"[{match.group(1)}]"
            run = para.add_run(citation_text)
            run.font.superscript = True
            run.font.name = self.FONT_SETTINGS["english_font"]
            run.font.color.rgb = RGBColor(0, 0, 0)

            last_end = match.end()

        if last_end < len(text):
            remaining_text = text[last_end:]
            run = para.add_run(remaining_text)
            run.font.name = self.FONT_SETTINGS["english_font"]
            run.font.color.rgb = RGBColor(0, 0, 0)
            if any("\u4e00" <= char <= "\u9fff" for char in remaining_text):
                self._set_east_asian_font(run, self.FONT_SETTINGS["chinese_font"])

    def _apply_mixed_fonts(self, para: "Paragraph") -> None:
        """为段落应用混合字体（中文宋体，英文Times New Roman）"""
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Pt, RGBColor
        except ImportError:
            return

        text = para.text
        if not text:
            return

        for run in para.runs:
            run.font.name = self.FONT_SETTINGS["english_font"]
            run.font.size = Pt(self.FONT_SIZES["body"])
            run.font.color.rgb = RGBColor(0, 0, 0)

            run_text = run.text
            if run_text and any("\u4e00" <= char <= "\u9fff" for char in run_text):
                self._set_east_asian_font(run, self.FONT_SETTINGS["chinese_font"])

    def format_references_section(
        self, paragraphs: List["Paragraph"], ctx: "ProcessingContext"
    ) -> None:
        """格式化参考文献章节

        将手打编号改为Word自动编号 [1], [2], [3]...

        Args:
            paragraphs: 段落列表
            ctx: 处理上下文
        """
        try:
            from docx.shared import Pt
        except ImportError:
            return

        ref_items = []
        ref_section_started = False

        for para in paragraphs:
            text = para.text.strip()
            text_no_space = text.replace(" ", "")

            # 使用更严格的匹配：必须是独立的标题，而不是包含关键词的正文
            is_ref_heading = text_no_space in ("参考文献", "references") or text.lower() in (
                "参考文献",
                "references",
            )
            if is_ref_heading:
                ref_section_started = True
                self._format_special_title(para)
                continue

            if ref_section_started and text and len(text) > 10:
                style_name = para.style.name.lower() if para.style and para.style.name else ""
                is_heading = bool(style_name.startswith("heading"))

                if not is_heading:
                    clean_text = re.sub(r"^\[\d+\]\s*", "", text)
                    clean_text = re.sub(r"^\d+[.．]\s*", "", clean_text)
                    ref_items.append((para, clean_text))

            if ref_section_started and text and len(text) < 10:
                if any(keyword in text for keyword in ["致谢", "附录", "结论"]):
                    break

        if ref_items:
            self._create_reference_numbering(ctx)

        for idx, (para, clean_text) in enumerate(ref_items, 1):
            para.clear()
            self._add_reference_text_with_mixed_fonts(para, clean_text)

            self._add_reference_bookmark(para, idx)

            para.paragraph_format.line_spacing = self.LINE_SPACING
            # 悬挂缩进由编号列表级别控制 (w:ind w:left="567" w:hanging="567")
            # 567 twips ≈ 1cm，与 [1] 编号的宽度匹配
            para.paragraph_format.space_after = Pt(3)

            self._apply_reference_numbering(para, idx)

    def _create_reference_numbering(self, ctx: "ProcessingContext") -> None:
        """创建参考文献的自定义编号格式 [1], [2], [3]..."""
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls

            try:
                numbering_part = ctx.doc.part.numbering_part
            except Exception:
                from docx.parts.numbering import NumberingPart

                numbering_part = NumberingPart.new()
                ctx.doc.part._package.parts.append(numbering_part)

            abstract_num_xml = parse_xml(r"""
                <w:abstractNum {} w:abstractNumId="99">
                    <w:multiLevelType w:val="singleLevel"/>
                    <w:lvl w:ilvl="0">
                        <w:start w:val="1"/>
                        <w:numFmt w:val="decimal"/>
                        <w:lvlText w:val="[%1]"/>
                        <w:lvlJc w:val="left"/>
                        <w:pPr>
                            <w:ind w:left="567" w:hanging="567"/>
                        </w:pPr>
                        <w:rPr>
                            <w:rFonts w:ascii="Times New Roman" w:eastAsia="Times New Roman"/>
                            <w:color w:val="000000"/>
                        </w:rPr>
                    </w:lvl>
                </w:abstractNum>
            """.format(nsdecls("w")))

            num_xml = parse_xml(r"""
                <w:num {} w:numId="99">
                    <w:abstractNumId w:val="99"/>
                    <w:lvlOverride w:ilvl="0">
                        <w:startOverride w:val="1"/>
                    </w:lvlOverride>
                </w:num>
            """.format(nsdecls("w")))

            numbering_element = numbering_part._element
            numbering_element.append(abstract_num_xml)
            numbering_element.append(num_xml)

        except Exception as e:
            ctx.add_warning(f"YZUThesisRule: 创建参考文献编号定义失败: {e}")

    def _apply_reference_numbering(self, paragraph, _number: int) -> None:
        """应用参考文献自动编号"""
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls

            pPr = paragraph._p.get_or_add_pPr()

            numPr = parse_xml(r"""
                <w:numPr {}>
                    <w:ilvl w:val="0"/>
                    <w:numId w:val="99"/>
                </w:numPr>
            """.format(nsdecls("w")))

            for child in list(pPr):
                if child.tag.endswith("numPr"):
                    pPr.remove(child)

            pPr.append(numPr)

        except Exception as e:
            import logging

            logging.debug(f"YZUThesisRule: 应用参考文献编号失败: {e}")

    def _add_reference_bookmark(self, para, ref_idx: int) -> None:
        """为参考文献条目添加书签，支持正文中的交叉引用

        书签命名规则：_RefRef{idx}（如 _RefRef1, _RefRef2）
        书签位置：段落起始处，覆盖整个条目内容

        Args:
            para: 参考文献段落
            ref_idx: 参考文献序号（从1开始）
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            bookmark_name = f"_RefRef{ref_idx}"
            bookmark_id = str(abs(hash(bookmark_name)) % 10000)

            bookmark_start = OxmlElement("w:bookmarkStart")
            bookmark_start.set(qn("w:id"), bookmark_id)
            bookmark_start.set(qn("w:name"), bookmark_name)

            para._p.insert(0, bookmark_start)

            bookmark_end = OxmlElement("w:bookmarkEnd")
            bookmark_end.set(qn("w:id"), bookmark_id)
            para._p.append(bookmark_end)

        except Exception as e:
            import logging

            logging.debug(f"YZUThesisRule: 添加参考文献书签失败 (ref={ref_idx}): {e}")

    def resolve_reference_cross_references(self, doc, ctx: "ProcessingContext") -> None:
        """解析正文中的文献引用上标，替换为REF域实现真正的交叉引用

        将正文中 [N] 形式的上标引用转换为 REF 域，
        指向文末对应参考文献的书签 (_RefRefN)。

        处理流程：
        1. 扫描正文段落中的 [N], [N-M], [N,M,...] 引用模式
        2. 对每个引用编号生成 REF 域指向 _RefRef{N}
        3. 保持上标格式不变

        Args:
            doc: 文档对象
            ctx: 处理上下文
        """
        citation_pattern = re.compile(r"\^?\[(\d+(?:[-,]\d+)*)\]\^?")

        for para in doc.paragraphs:
            text = para.text
            if not text:
                continue

            para_type = self.detect_paragraph_type(para, ctx)
            if para_type in (
                "references",
                "caption",
                "title",
                "heading",
                "special_title",
                "signature",
            ):
                continue

            matches = list(citation_pattern.finditer(text))
            if not matches:
                continue

            para.clear()

            last_end = 0
            for match in matches:
                if match.start() > last_end:
                    normal_text = text[last_end : match.start()]
                    self._add_text_run(
                        para, normal_text, use_chinese_font=self._has_chinese(normal_text)
                    )

                citation_raw = match.group(1)
                ref_numbers = self._parse_citation_numbers(citation_raw)

                for i, ref_num in enumerate(ref_numbers):
                    if i > 0:
                        self._add_text_run(para, ",", is_superscript=True)

                    bookmark_name = f"_RefRef{ref_num}"
                    self._build_ref_field(para, bookmark_name, f"[{ref_num}]", is_superscript=True)

                last_end = match.end()

            if last_end < len(text):
                remaining_text = text[last_end:]
                # 尝试保留原有 run 的属性（如果不复杂），或者简单添加
                self._add_text_run(
                    para, remaining_text, use_chinese_font=self._has_chinese(remaining_text)
                )

    def _add_text_run(
        self, para, text: str, use_chinese_font: bool = False, is_superscript: bool = False
    ) -> None:
        """辅助方法：添加文本运行并设置基本字体"""
        from docx.shared import Pt, RGBColor

        run = para.add_run(text)
        run.font.name = self.FONT_SETTINGS["english_font"]
        run.font.size = Pt(self.FONT_SIZES["body"])
        if is_superscript:
            run.font.superscript = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        if use_chinese_font:
            self._set_east_asian_font(run, self.FONT_SETTINGS["chinese_font"])

    def _has_chinese(self, text: str) -> bool:
        """检查及是否包含中文字段"""
        return any("\u4e00" <= char <= "\u9fff" for char in text)

    def _parse_citation_numbers(self, citation_raw: str) -> List[int]:
        """解析引用编号，支持 [1-3] 或 [1,2,3]"""
        nums = []
        parts = citation_raw.split(",")
        for part in parts:
            if "-" in part:
                try:
                    start, end = map(int, part.split("-"))
                    nums.extend(range(start, end + 1))
                except ValueError:
                    pass
            else:
                try:
                    nums.append(int(part))
                except ValueError:
                    pass
        return nums

    def _build_ref_field(
        self, para, bookmark_name: str, display_text: str, is_superscript: bool = False
    ) -> None:
        """构建 REF 交叉引用域"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        run_begin = para.add_run()
        if is_superscript:
            run_begin.font.superscript = True
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run_begin._r.append(fldChar1)

        run_instr = para.add_run()
        if is_superscript:
            run_instr.font.superscript = True
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" REF {bookmark_name} \\h "
        run_instr._r.append(instr)

        run_sep = para.add_run()
        if is_superscript:
            run_sep.font.superscript = True
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        run_sep._r.append(fldChar2)

        run_text = para.add_run(display_text)
        run_text.font.name = self.FONT_SETTINGS["english_font"]
        run_text.font.size = Pt(self.FONT_SIZES["body"])
        if is_superscript:
            run_text.font.superscript = True

        run_end = para.add_run()
        if is_superscript:
            run_end.font.superscript = True
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run_end._r.append(fldChar3)

    def _add_reference_text_with_mixed_fonts(self, paragraph, text: str) -> None:
        """为参考文献添加文本，中英文使用不同字体"""
        try:
            from docx.shared import Pt, RGBColor
        except ImportError:
            paragraph.add_run(text)
            return

        current_text = ""
        is_current_chinese = None

        for char in text:
            is_chinese = "\u4e00" <= char <= "\u9fff"

            if is_current_chinese is None:
                is_current_chinese = is_chinese
                current_text = char
            elif is_current_chinese == is_chinese:
                current_text += char
            else:
                if current_text:
                    run = paragraph.add_run(current_text)
                    run.font.size = Pt(self.FONT_SIZES["body"])
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    if is_current_chinese:
                        run.font.name = self.FONT_SETTINGS["chinese_font"]
                        self._set_east_asian_font(run, self.FONT_SETTINGS["chinese_font"])
                    else:
                        run.font.name = self.FONT_SETTINGS["english_font"]
                is_current_chinese = is_chinese
                current_text = char

        if current_text:
            run = paragraph.add_run(current_text)
            run.font.size = Pt(self.FONT_SIZES["body"])
            run.font.color.rgb = RGBColor(0, 0, 0)
            if is_current_chinese:
                run.font.name = self.FONT_SETTINGS["chinese_font"]
                self._set_east_asian_font(run, self.FONT_SETTINGS["chinese_font"])
            else:
                run.font.name = self.FONT_SETTINGS["english_font"]

    def _deduplicate_overlapping_refs(self, all_refs: list) -> list:
        """去除重叠的交叉引用匹配

        当多个正则模式匹配到同一段文本的重叠区域时（如
        "根据公式(1-1)" 同时被带前缀模式和无前缀模式匹配），
        只保留匹配范围最大的那个，避免重复输出。

        策略：按起始位置排序后遍历，若当前匹配与已保留匹配
        有交集则跳过（优先保留先出现的/更长的匹配）。

        Args:
            all_refs: [(类型, match对象), ...] 列表

        Returns:
            去重后的列表
        """
        if len(all_refs) <= 1:
            return all_refs

        sorted_refs = sorted(all_refs, key=lambda x: (x[1].start(), -(x[1].end() - x[1].start())))
        kept = []

        for ref in sorted_refs:
            start, end = ref[1].start(), ref[1].end()
            overlaps = False
            for kept_ref in kept:
                ks, ke = kept_ref[1].start(), kept_ref[1].end()
                if start < ke and end > ks:
                    overlaps = True
                    break
            if not overlaps:
                kept.append(ref)

        return kept

    def resolve_cross_references(self, doc, ctx: "ProcessingContext") -> None:
        """解析正文中的图表和公式交叉引用，使用REF域实现真正的交叉引用

        优化版本：合并公式和图表引用解析为单次遍历，避免重复段落重建。

        支持的引用模式：
        - 图表: "图X-Y" / "表X-Y" / "如图X-Y" / "见表X-Y"
        - 公式: "公式(X-Y)" / "式(X-Y)" / "根据公式(X-Y)"

        输出格式：
        - 图表: "图 X-Y" / "表 X-Y"（REF域placeholder包含完整格式）
        - 公式: "(X-Y)"（REF域placeholder包含完整格式）

        Args:
            doc: 文档对象
            ctx: 处理上下文
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Pt, RGBColor
        except ImportError:
            return

        fig_nums = self._collect_caption_numbers(doc, "图")
        tab_nums = self._collect_caption_numbers(doc, "表")
        formula_nums = self._collect_formula_numbers(doc)

        if not fig_nums and not tab_nums and not formula_nums:
            return

        formula_patterns = [
            re.compile(
                r"((?:根据|由|如|见|参|用|代入|利用)\s*(?:公式|式))\s*(\(\d+(?:[-‐–—]\d+)\))"
            ),
            re.compile(r"(公式|式)\s*(\(\d+(?:[-‐–—]\d+)\))"),
        ]

        fig_table_patterns = [
            re.compile(r"((?:如|见|参|看|详|示)?\s*(图|表))\s*(\d+[-‐–—]\d+)"),
        ]

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or self._is_caption_text(para, text):
                continue

            para_type = self.detect_paragraph_type(para, ctx)
            if para_type in ("caption", "title", "heading", "special_title", "signature", "empty"):
                continue

            original_text = para.text

            all_refs = []

            for pattern in formula_patterns:
                for match in pattern.finditer(original_text):
                    all_refs.append(("formula", match))

            for pattern in fig_table_patterns:
                for match in pattern.finditer(original_text):
                    all_refs.append(("figtable", match))

            if not all_refs:
                continue

            all_refs = self._deduplicate_overlapping_refs(all_refs)

            all_refs.sort(key=lambda x: x[1].start(), reverse=True)

            self._rebuild_paragraph_with_refs(
                para, original_text, all_refs, formula_nums, fig_nums, tab_nums
            )

    def _rebuild_paragraph_with_refs(
        self, para, text: str, refs: list, formula_nums: list, fig_nums: list, tab_nums: list
    ) -> None:
        """单次重建段落，处理所有引用

        Args:
            para: 段落对象
            text: 原始文本
            refs: 引用列表 [(类型, match), ...]
            formula_nums: 公式编号列表
            fig_nums: 图编号列表
            tab_nums: 表编号列表
        """
        para.clear()

        last_end = 0

        refs_sorted = sorted(refs, key=lambda x: x[1].start())

        for ref_type, match in refs_sorted:
            start = match.start()
            end = match.end()

            if start > last_end:
                before_text = text[last_end:start]
                self._add_text_run(
                    para, before_text, use_chinese_font=self._has_chinese(before_text)
                )

            if ref_type == "formula":
                self._process_formula_ref(para, match, formula_nums)
            else:
                self._process_figtable_ref(para, match, fig_nums, tab_nums)

            last_end = end

        if last_end < len(text):
            after_text = text[last_end:]
            self._add_text_run(para, after_text, use_chinese_font=self._has_chinese(after_text))

    def _process_formula_ref(self, para, match, formula_nums: list) -> None:
        """处理单个公式引用"""
        prefix_text = match.group(1) if len(match.groups()) >= 1 and match.group(1) else ""
        old_ref = match.group(2) if len(match.groups()) >= 2 else ""

        if not old_ref:
            return

        ref_clean = old_ref.strip("()").replace("–", "-").replace("—", "-").replace("‐", "-")
        replacement_num = self._find_replacement_number(ref_clean, formula_nums)
        if not replacement_num:
            replacement_num = ref_clean

        bookmark_name = f"_RefFormula{replacement_num}"
        full_ref = f"({replacement_num})"

        if prefix_text:
            self._add_text_run(para, prefix_text, use_chinese_font=self._has_chinese(prefix_text))

        self._build_ref_field(para, bookmark_name, full_ref)

    def _process_figtable_ref(self, para, match, fig_nums: list, tab_nums: list) -> None:
        """处理单个图表引用"""
        prefix_and_label = match.group(1) or ""
        label = match.group(2) or ""
        old_ref = match.group(3) or ""

        if not old_ref or not label:
            return

        prefix_text = prefix_and_label[: -len(label)] if len(prefix_and_label) > len(label) else ""

        label_key = "图" if label and label[0] in "图Fig" else "表"
        nums_list = fig_nums if label_key == "图" else tab_nums

        if not nums_list:
            return

        replacement_num = self._find_replacement_number(old_ref, nums_list)
        if not replacement_num:
            return

        bookmark_name = f"_Ref{label}{replacement_num}"
        full_display = f"{label} {replacement_num}"

        if prefix_text and not prefix_text.endswith(label):
            self._add_text_run(
                para, prefix_text + " ", use_chinese_font=self._has_chinese(prefix_text)
            )

        self._build_ref_field(para, bookmark_name, full_display)

    def _build_caption_map(self, doc) -> dict:
        """构建题注编号映射 {类型: {原始引用 → 实际编号}}

        扫描文档中所有题注段落，提取实际编号。
        同时构建反向映射用于交叉引用替换。

        Returns:
            {'图': {原始引用: 实际编号}, '表': {...}}
        """
        caption_map = {"图": {}, "表": {}}
        all_fig_nums = []
        all_tab_nums = []

        fig_pattern = re.compile(r"^图\s+(\d+[-‐–—]\d+)", re.IGNORECASE)
        tab_pattern = re.compile(r"^表\s+(\d+[-‐–—]\d+)", re.IGNORECASE)

        for para in doc.paragraphs:
            text = para.text.strip()

            fm = fig_pattern.match(text)
            if fm:
                all_fig_nums.append(fm.group(1))
                continue

            tm = tab_pattern.match(text)
            if tm:
                all_tab_nums.append(tm.group(1))
                continue

        for i, num in enumerate(all_fig_nums):
            if i > 0:
                prev_parts = all_fig_nums[i - 1].split("-")
                curr_parts = num.split("-")
                if len(prev_parts) == 2 and len(curr_parts) == 2:
                    if curr_parts[0] != prev_parts[0]:
                        expected = f"{curr_parts[0]}-1"
                    else:
                        expected = f"{curr_parts[0]}-{int(curr_parts[1])}"
                else:
                    expected = str(i + 1)
            else:
                expected = num.split("-")[0] + "-1" if "-" in num else "1"
            caption_map["图"][expected] = num

        for i, num in enumerate(all_tab_nums):
            if i > 0:
                prev_parts = all_tab_nums[i - 1].split("-")
                curr_parts = num.split("-")
                if len(prev_parts) == 2 and len(curr_parts) == 2:
                    if curr_parts[0] != prev_parts[0]:
                        expected = f"{curr_parts[0]}-1"
                    else:
                        expected = f"{curr_parts[0]}-{int(curr_parts[1])}"
                else:
                    expected = str(i + 1)
            else:
                expected = num.split("-")[0] + "-1" if "-" in num else "1"
            caption_map["表"][expected] = num

        for num in all_fig_nums:
            caption_map["图"][num] = num

        for num in all_tab_nums:
            caption_map["表"][num] = num

        return caption_map

    def _is_caption_text(self, para, text: str) -> bool:
        """检测是否为题注段落文本"""
        style_name = para.style.name.lower() if para.style and para.style.name else ""
        if style_name in ("caption",):
            return True
        if re.match(r"^(图|表|Figure|Table)\s*\d+", text, re.IGNORECASE):
            return True
        return False

    def format_tables(self, doc) -> None:
        """格式化文档中的所有表格（三线表格式）

        三线表结构（学术论文标准）：
        ┌─────────────────────────────┐  ← 顶线（粗线 1.5pt）
        │ 表头  │ 表头  │ 表头       │
        ├─────────────────────────────┤  ← 栏目线（细线 0.5pt）
        │ 数据  │ 数据  │ 数据       │
        │ 数据  │ 数据  │ 数据       │
        └─────────────────────────────┘  ← 底线（粗线 1.5pt）

        特点：
        - 仅保留三条横线，无竖线、无内部横线
        - 顶线和底线使用粗线（1.5pt）强调表格边界
        - 栏目线使用细线（0.5pt）分隔表头与数据区

        Args:
            doc: 文档对象
        """
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt, RGBColor
        except ImportError:
            return

        for table in doc.tables:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    for para in cell.paragraphs:
                        para.paragraph_format.space_before = Pt(3)
                        para.paragraph_format.space_after = Pt(3)

                        for run in para.runs:
                            run.font.size = Pt(self.FONT_SIZES["body"])
                            run.font.name = self.FONT_SETTINGS["english_font"]
                            run.font.color.rgb = RGBColor(0, 0, 0)

                            run_text = run.text
                            if run_text and any("\u4e00" <= char <= "\u9fff" for char in run_text):
                                self._set_east_asian_font(run, self.FONT_SETTINGS["chinese_font"])

            self._set_three_line_table_borders(table)

    def _set_three_line_table_borders(self, table) -> None:
        """设置三线表边框（学术论文标准格式）- 使用单元格级边框

        三线表由三条横线组成：
        - 顶线（top）：粗线 1.5pt，表格顶部边界
        - 栏目线（header bottom）：细线 0.5pt，表头行底部
        - 底线（bottom）：粗线 1.5pt，表格底部边界

        无竖线、无内部横线、无左右边框。

        实现策略：
        使用单元格级边框（tcBorders），比表格级边框更可靠：
        1. 移除 tblStyle 和 tblLook，防止 Word 应用默认表格样式
        2. 首行单元格：top=粗线(1.5pt) + bottom=细线(0.5pt)
        3. 中间行单元格：无边框
        4. 末行单元格：bottom=粗线(1.5pt)
        5. 所有单元格：left=nil, right=nil（无竖线）

        Args:
            table: 表格对象
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ImportError:
            return

        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

        # 移除 tblStyle，防止 Word 应用默认表格样式（含竖线）
        existing_style = tblPr.find(qn("w:tblStyle"))
        if existing_style is not None:
            tblPr.remove(existing_style)

        # 完全移除 tblLook，防止 Word 应用任何默认样式
        existing_look = tblPr.find(qn("w:tblLook"))
        if existing_look is not None:
            tblPr.remove(existing_look)

        # 设置表格级边框为 nil，确保不显示任何默认边框
        borders = OxmlElement("w:tblBorders")
        for border_name in ["top", "bottom", "left", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            borders.append(border)

        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(borders)

        # 使用单元格级边框设置三线表
        self._set_cell_level_three_line_borders(table)

    def _set_cell_level_three_line_borders(self, table) -> None:
        """使用单元格级边框设置三线表格式

        单元格级边框优先级高于表格级边框，更可靠。

        三线表结构：
        - 首行：top=粗线(1.5pt) + bottom=细线(0.5pt)  → 顶线 + 栏目线
        - 中间行：无边框
        - 末行：bottom=粗线(1.5pt)  → 底线
        - 所有单元格：left=nil, right=nil  → 无竖线

        Args:
            table: 表格对象
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ImportError:
            return

        if len(table.rows) == 0:
            return

        last_row_idx = len(table.rows) - 1

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.tcPr if tc.tcPr is not None else OxmlElement("w:tcPr")
                if tc.tcPr is None:
                    tc.insert(0, tcPr)

                tcBorders = OxmlElement("w:tcBorders")

                if row_idx == 0:
                    # 首行：顶线（粗线 1.5pt）
                    top = OxmlElement("w:top")
                    top.set(qn("w:val"), "single")
                    top.set(qn("w:sz"), "12")  # 1.5pt
                    top.set(qn("w:space"), "0")
                    top.set(qn("w:color"), "000000")
                    tcBorders.append(top)

                    # 栏目线（细线 0.75pt）
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "6")  # 0.75pt
                    bottom.set(qn("w:space"), "0")
                    bottom.set(qn("w:color"), "000000")
                    tcBorders.append(bottom)
                elif row_idx == last_row_idx:
                    # 末行：底线（粗线 1.5pt）
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "12")  # 1.5pt
                    bottom.set(qn("w:space"), "0")
                    bottom.set(qn("w:color"), "000000")
                    tcBorders.append(bottom)

                    # 其他边框为无
                    top = OxmlElement("w:top")
                    top.set(qn("w:val"), "nil")
                    tcBorders.append(top)
                else:
                    # 中间行：无边框
                    top = OxmlElement("w:top")
                    top.set(qn("w:val"), "nil")
                    tcBorders.append(top)

                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "nil")
                    tcBorders.append(bottom)

                # 所有单元格：无竖线
                left = OxmlElement("w:left")
                left.set(qn("w:val"), "nil")
                tcBorders.append(left)

                right = OxmlElement("w:right")
                right.set(qn("w:val"), "nil")
                tcBorders.append(right)

                existing = tcPr.find(qn("w:tcBorders"))
                if existing is not None:
                    tcPr.remove(existing)
                tcPr.append(tcBorders)

    def _set_header_row_bottom_border(self, table) -> None:
        """为表头行（第一行）的每个单元格设置底部细线（栏目线）

        这是三线表的第二条线，位于表头与数据区之间。
        使用 0.5pt 细线以区别于顶线和底线的 1.5pt 粗线。

        Args:
            table: 表格对象
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ImportError:
            return

        if len(table.rows) == 0:
            return

        header_row = table.rows[0]

        for cell in header_row.cells:
            tc = cell._tc
            tcPr = tc.tcPr if tc.tcPr is not None else OxmlElement("w:tcPr")
            if tc.tcPr is None:
                tc.insert(0, tcPr)

            tcBorders = OxmlElement("w:tcBorders")

            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "000000")
            tcBorders.append(bottom)

            for edge in ["top", "left", "right"]:
                edge_elem = OxmlElement(f"w:{edge}")
                edge_elem.set(qn("w:val"), "nil")
                tcBorders.append(edge_elem)

            existing = tcPr.find(qn("w:tcBorders"))
            if existing is not None:
                tcPr.remove(existing)
            tcPr.append(tcBorders)

    def _set_table_borders(self, table) -> None:
        """设置表格边框（单线0.5pt黑色）—— 已弃用，保留向后兼容"""

    def format_formulas(self, doc) -> None:
        """格式化文档中的数学公式（OMML）

        区分两种公式类型：
        1. 行内公式：嵌入在正文段落中（如 $O(n)$），保持左对齐融入正文
        2. 独立公式：Pandoc用<m:oMathPara>包裹，单独成段，需要添加编号

        Pandoc输出结构：
        - 独立公式: <w:p><m:oMathPara><m:oMath>...</m:oMath></m:oMathPara></w:p>
          注意：这种段落para.runs为空！内容在oMathPara中
        - 行内公式: <w:p><w:r>...<m:oMath>...</m:oMath><w:r>...</w:p>

        编号规则：
        - 使用章节化编号格式（章节号-序号），如 (1-1), (3-2)
        - 在每个一级标题下独立计数

        Args:
            doc: 文档对象
        """
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt, RGBColor
        except ImportError:
            return

        formula_chapters = self._scan_formula_chapters(doc)
        chapter_formula_counters: Dict[int, int] = {}

        display_formula_indices = []

        for para_idx, para in enumerate(doc.paragraphs):
            try:
                xml = para._p.xml
                if "m:oMathPara" in xml:
                    display_formula_indices.append(para_idx)
            except Exception:
                pass

        removed_offset = 0
        for para_idx in display_formula_indices:
            adjusted_idx = para_idx - removed_offset
            if adjusted_idx >= len(doc.paragraphs):
                continue

            para = doc.paragraphs[adjusted_idx]

            chapter_num = formula_chapters.get(para_idx, 1)
            if chapter_num not in chapter_formula_counters:
                chapter_formula_counters[chapter_num] = 0
            chapter_formula_counters[chapter_num] += 1
            formula_seq = chapter_formula_counters[chapter_num]

            removed = self._format_display_formula_v3(doc, adjusted_idx, chapter_num, formula_seq)
            if removed != adjusted_idx:
                removed_offset += 1

        for para_idx, para in enumerate(doc.paragraphs):
            try:
                xml = para._p.xml
                has_math = "m:oMath" in xml or "oMath" in xml
                has_omath_para = "m:oMathPara" in xml
            except Exception:
                continue

            if not has_math or has_omath_para:
                continue

            is_inline = self._is_inline_formula(para)
            if is_inline:
                self._format_inline_formula(para)

    def _format_display_formula_v3(
        self, doc, para_idx: int, chapter_num: int, formula_seq: int
    ) -> int:
        """公式格式化方法（V4）：将oMathPara转为行内oMath+Tab+编号

        Word兼容性核心发现：
        - <m:oMathPara> 是一个自包含的块元素，会吞掉外部插入的制表符
        - WPS 能正确处理这种情况，但 Word 不行
        - 解决方案：提取 <m:oMath> 内容，移除 <m:oMathPara> 包裹，
          使其成为段落内的行内元素，然后 Tab 就能正常工作

        最终段落结构：
        <w:p>
          <w:bookmarkStart/>
          <w:r><w:t>\t</t></w:r>           ← 居中制表符
          <m:oMath>...</m:oMath>             ← 行内公式（不再是oMathPara！）
          <w:r><w:t>\t(1-1)</w:t></w:r>    ← 右对齐制表符+编号
          <w:bookmarkEnd/>
        </w:p>

        Args:
            doc: 文档对象
            para_idx: 公式段落在文档中的索引
            chapter_num: 章节号
            formula_seq: 该章节内的公式序号

        Returns:
            被移除的编号段落的索引（如果有），否则返回para_idx
        """
        import logging

        logger = logging.getLogger(__name__)

        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt
        except ImportError:
            return para_idx

        equation_number = f"({chapter_num}-{formula_seq})"
        bookmark_name = f"_RefFormula{chapter_num}-{formula_seq}"
        bookmark_id = str(abs(hash(bookmark_name)) % 10000)

        para = doc.paragraphs[para_idx]
        removed_idx = self._remove_following_number_paragraph(doc, para_idx, logger)

        try:
            self._convert_omathpara_to_inline(para)
            self._setup_formula_tab_stops(para)
            self._set_formula_paragraph_format(para)
            self._add_formula_bookmark_and_number(para, bookmark_id, bookmark_name, equation_number)

            logger.debug(f"公式V4格式化完成: {equation_number}, 移除编号段落索引={removed_idx}")

        except Exception as e:
            logger.warning(f"公式V4格式化失败: {e}，使用降级模式")
            import traceback

            logger.debug(traceback.format_exc())
            try:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = None
            except Exception:
                pass

        return removed_idx

    def _convert_omathpara_to_inline(self, para) -> None:
        """将 oMathPara 转换为行内 oMath 元素

        Args:
            para: 段落对象
        """
        M_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

        omath_para = para._p.find(f".//{M_NS}oMathPara")
        if omath_para is None:
            return

        omath_elements = []
        for child in list(omath_para):
            tag = child.tag
            if tag == f"{M_NS}oMathParaPr":
                continue
            if tag == f"{M_NS}oMath":
                omath_elements.append(child)

        para._p.remove(omath_para)

        for elem in omath_elements:
            para._p.append(elem)

        self._set_omath_font(para)

    def _set_omath_font(self, para) -> None:
        """为 OMML 公式元素设置 Times New Roman 字体

        OMML (Office Math Markup Language) 公式元素使用特殊的字体设置方式。
        需要在 m:oMath 元素中递归查找所有 m:r (文本运行) 元素，
        并为其设置 w:rFonts 属性以确保使用 Times New Roman 字体。

        Args:
            para: 包含 OMML 公式的段落对象
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ImportError:
            return

        M_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
        W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

        font_name = self.FONT_SETTINGS["english_font"]

        omath_elements = para._p.findall(f".//{M_NS}oMath")
        for omath in omath_elements:
            self._apply_font_to_omath_element(omath, font_name, M_NS, W_NS)

    def _apply_font_to_omath_element(self, element, font_name: str, M_NS: str, W_NS: str) -> None:
        """递归为 OMML 元素中的所有文本节点应用字体设置

        Args:
            element: OMML XML 元素
            font_name: 目标字体名称（如 'Times New Roman'）
            M_NS: Math 命名空间
            W_NS: Word 命名空间
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        for child in element:
            tag = child.tag

            if tag == f"{M_NS}r":
                rPr = child.find(f"{W_NS}rPr")
                if rPr is None:
                    rPr_idx = 0
                    for idx, sibling in enumerate(child):
                        sibling_tag = getattr(sibling, "tag", "")
                        if (
                            isinstance(sibling_tag, str)
                            and sibling_tag.startswith(W_NS)
                            and "Pr" in sibling_tag
                        ):
                            rPr_idx = idx + 1
                            break
                    rPr = OxmlElement("w:rPr")
                    child.insert(rPr_idx, rPr)

                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.insert(0, rFonts)

                rFonts.set(qn("w:ascii"), font_name)
                rFonts.set(qn("w:hAnsi"), font_name)
            else:
                self._apply_font_to_omath_element(child, font_name, M_NS, W_NS)

    def _setup_formula_tab_stops(self, para) -> None:
        """设置公式段落的制表位（居中+右对齐）

        Args:
            para: 段落对象
        """
        try:
            from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm
        except ImportError:
            return

        pPr = para._p.get_or_add_pPr()

        existing_tabs = pPr.find(qn("w:tabs"))
        if existing_tabs is not None:
            pPr.remove(existing_tabs)

        tabs = OxmlElement("w:tabs")

        tab_center = OxmlElement("w:tab")
        tab_center.set(qn("w:val"), "center")
        tab_center.set(qn("w:pos"), str(int(Cm(self.TAB_CENTER_POS_CM).twips)))
        tab_center.set(qn("w:leader"), "spaces")
        tabs.append(tab_center)

        tab_right = OxmlElement("w:tab")
        tab_right.set(qn("w:val"), "right")
        tab_right.set(qn("w:pos"), str(int(Cm(self.TAB_RIGHT_POS_CM).twips)))
        tab_right.set(qn("w:leader"), "spaces")
        tabs.append(tab_right)

        pPr.insert(0, tabs)

        pf = para.paragraph_format
        while len(pf.tab_stops) > 0:
            try:
                pf.tab_stops._element.remove(pf.tab_stops._element[0])
            except Exception:
                break
        pf.tab_stops.add_tab_stop(
            Cm(self.TAB_CENTER_POS_CM), WD_TAB_ALIGNMENT.CENTER, WD_TAB_LEADER.SPACES
        )
        pf.tab_stops.add_tab_stop(
            Cm(self.TAB_RIGHT_POS_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
        )

    def _set_formula_paragraph_format(self, para) -> None:
        """设置公式段落的格式属性

        Args:
            para: 段落对象
        """
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Cm, Pt
        except ImportError:
            return

        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Cm(0)

    def _add_formula_bookmark_and_number(
        self, para, bookmark_id: str, bookmark_name: str, equation_number: str
    ) -> None:
        """添加公式书签和编号

        Args:
            para: 段落对象
            bookmark_id: 书签ID
            bookmark_name: 书签名称
            equation_number: 公式编号
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Pt
        except ImportError:
            return

        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), bookmark_id)
        bookmark_start.set(qn("w:name"), bookmark_name)
        para._p.insert(0, bookmark_start)

        tab_run = OxmlElement("w:r")
        tab_rpr = OxmlElement("w:rPr")
        tab_fonts = OxmlElement("w:rFonts")
        tab_fonts.set(qn("w:ascii"), self.FONT_SETTINGS["english_font"])
        tab_fonts.set(qn("w:hAnsi"), self.FONT_SETTINGS["english_font"])
        tab_rpr.append(tab_fonts)
        tab_sz = OxmlElement("w:sz")
        tab_sz.set(qn("w:val"), str(int(Pt(self.FONT_SIZES["body"]).pt * 2)))
        tab_rpr.append(tab_sz)
        tab_run.append(tab_rpr)
        tab_text = OxmlElement("w:t")
        tab_text.text = "\t"
        tab_text.set(qn("xml:space"), "preserve")
        tab_run.append(tab_text)
        para._p.insert(1, tab_run)

        num_run = OxmlElement("w:r")
        num_rpr = OxmlElement("w:rPr")
        num_fonts = OxmlElement("w:rFonts")
        num_fonts.set(qn("w:ascii"), self.FONT_SETTINGS["english_font"])
        num_fonts.set(qn("w:hAnsi"), self.FONT_SETTINGS["english_font"])
        num_rpr.append(num_fonts)
        num_sz = OxmlElement("w:sz")
        num_sz.set(qn("w:val"), str(int(Pt(self.FONT_SIZES["body"]).pt * 2)))
        num_rpr.append(num_sz)
        num_color = OxmlElement("w:color")
        num_color.set(qn("w:val"), "000000")
        num_rpr.append(num_color)
        num_run.append(num_rpr)
        num_text_elem = OxmlElement("w:t")
        num_text_elem.text = "\t" + equation_number
        num_run.append(num_text_elem)
        para._p.append(num_run)

        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), bookmark_id)
        para._p.append(bookmark_end)

    def _remove_following_number_paragraph(self, doc, formula_para_idx: int, logger) -> int:
        """检测并移除紧跟在公式段落后的编号段落

        Pandoc可能将编号（如 (1-1) 或 （1-1））输出为单独的段落。
        这个方法检测并移除该段落以避免双编号。

        Args:
            doc: 文档对象
            formula_para_idx: 公式段落的索引
            logger: 日志记录器

        Returns:
            被移除的段落索引，如果没有移除则返回formula_para_idx
        """
        import re

        number_patterns = [
            re.compile(r"^[\s\(（]*\d+[-‐–—]\d+[\s\)）]*$"),
            re.compile(r"^[\s\(（]?\d+[-‐–—]\d+[\s\)）]?$"),
        ]

        next_idx = formula_para_idx + 1
        if next_idx >= len(doc.paragraphs):
            return formula_para_idx

        next_para = doc.paragraphs[next_idx]
        text = next_para.text.strip()

        if not text:
            return formula_para_idx

        for pattern in number_patterns:
            if pattern.match(text):
                try:
                    parent = next_para._p.getparent()
                    parent.remove(next_para._p)
                    logger.debug(f"  移除编号段落 #{next_idx}: '{text}'")
                    return next_idx
                except Exception as e:
                    logger.debug(f"  移除编号段落失败: {e}")
                    break

        return formula_para_idx

    def _scan_formula_chapters(self, doc) -> Dict[int, int]:
        """扫描文档，确定每个段落所属的章节号（用于公式编号）

        从一级标题文本中提取章节号（支持手打标题格式如 "4. 实验结果"）
        - 跳过第一个 H1（论文标题）
        - 跳过非章节特殊标题（摘要、参考文献等）

        Args:
            doc: 文档对象

        Returns:
            字典: {段落索引: 章节号}
        """
        paragraph_chapters: Dict[int, int] = {}
        current_chapter = 1
        first_h1_seen = False

        non_chapter_keywords = {
            "摘要",
            "abstract",
            "关键词",
            "keywords",
            "参考文献",
            "references",
            "目录",
            "contents",
            "致谢",
            "acknowledgements",
            "附录",
            "appendix",
            "结论",
            "conclusion",
            "引言",
            "introduction",
        }

        # 匹配手打章节号的正则：支持 "1."、"第一章"、"第1章" 等格式
        chapter_pattern = re.compile(r"^(\d+)[\.\s]")
        chinese_chapter_pattern = re.compile(r"第([一二三四五六七八九十\d]+)章")

        for i, para in enumerate(doc.paragraphs):
            if self._is_heading_level_1_for_chapter(para):
                if not first_h1_seen:
                    first_h1_seen = True
                    paragraph_chapters[i] = current_chapter
                    continue

                text = para.text.strip()
                text_lower = text.lower().replace(" ", "")
                is_non_chapter = any(kw in text_lower for kw in non_chapter_keywords)

                if not is_non_chapter:
                    # 尝试从标题文本中提取章节号
                    chapter_num = self._extract_chapter_number(text)
                    if chapter_num > 0:
                        current_chapter = chapter_num
                    else:
                        current_chapter += 1

            paragraph_chapters[i] = current_chapter

        return paragraph_chapters

    def _extract_chapter_number(self, text: str) -> int:
        """从标题文本中提取章节号

        支持格式：
        - "1. 绪论" -> 1
        - "4. 实验结果" -> 4
        - "第一章 绪论" -> 1
        - "第1章 绪论" -> 1

        Args:
            text: 标题文本

        Returns:
            章节号，如果无法提取则返回 0
        """
        import re

        # 匹配 "1."、"4." 等阿拉伯数字格式
        match = re.match(r"^(\d+)[\.\s]", text.strip())
        if match:
            return int(match.group(1))

        # 匹配 "第一章"、"第1章" 等中文格式
        match = re.match(r"第([一二三四五六七八九十\d]+)章", text.strip())
        if match:
            num_str = match.group(1)
            # 转换中文数字
            chinese_nums = {
                "一": 1,
                "二": 2,
                "三": 3,
                "四": 4,
                "五": 5,
                "六": 6,
                "七": 7,
                "八": 8,
                "九": 9,
                "十": 10,
            }
            if num_str in chinese_nums:
                return chinese_nums[num_str]
            elif num_str.isdigit():
                return int(num_str)

        return 0

    def _is_heading_level_1_for_chapter(self, para) -> bool:
        """检测是否为一级标题（用于章节号计算）"""
        if not para.style:
            return False
        style_name = para.style.name.lower()
        if not style_name.startswith("heading"):
            return False
        match = re.match(r"heading\s*(\d+)", style_name)
        if match and match.group(1) == "1":
            return True
        return "heading 1" in style_name or style_name == "heading1"

    def _format_inline_formula(self, para) -> None:
        """格式化行内公式：保持左对齐，融入正文流

        行内公式不改变段落对齐方式，仅统一字体样式，
        确保与周围正文的视觉一致性。
        """
        try:
            from docx.shared import Pt, RGBColor
        except ImportError:
            return

        for run in para.runs:
            run.font.name = self.FONT_SETTINGS["english_font"]
            run.font.size = Pt(self.FONT_SIZES["body"])
            run.font.color.rgb = RGBColor(0, 0, 0)

        self._set_omath_font(para)

    def _format_display_formula(self, para, chapter_num: int = 1, formula_seq: int = 1) -> None:
        """格式化独立公式：公式居中，编号右对齐（Word兼容版本）

        布局结构：
        ┌─────────────────────────────────────────────┐
        │              公式内容              (1-1)    │
        │          [公式居中]          [编号右对齐]   │
        └─────────────────────────────────────────────┘

        编号格式：使用章节化编号（章节号-序号），如 (1-1), (3-2)
        书签支持：添加 _RefFormula{chapter_num}-{seq} 书签用于交叉引用

        使用 Tab 制表位实现：
        - 居中制表位：公式位置（页面中心）
        - 右制表位：编号位置（右侧）

        Word兼容性改进：
        - 先移除原始文档中已有的编号文本（支持中英文括号，避免双编号）
        - 在正确位置插入书签、制表符和编号
        - 保留原有 OMML 公式元素的命名空间完整性

        Args:
            para: 段落对象
            chapter_num: 章节号
            formula_seq: 该章节内的公式序号
        """
        import logging

        logger = logging.getLogger(__name__)

        try:
            import re

            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
            from docx.oxml import OxmlElement, parse_xml
            from docx.oxml.ns import qn
            from docx.shared import Cm, Emu, Pt, RGBColor, Twips
        except ImportError:
            return

        equation_number = f"({chapter_num}-{formula_seq})"
        bookmark_name = f"_RefFormula{chapter_num}-{formula_seq}"
        bookmark_id = str(abs(hash(bookmark_name)) % 10000)

        try:
            self._remove_existing_formula_number(para, logger)

            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = Cm(0)

            self._setup_formula_tab_stops(para)

            logger.debug(f"_format_display_formula: 开始处理公式段落，编号={equation_number}")

            self._format_display_formula_safe(
                para, equation_number, bookmark_name, bookmark_id, logger
            )

        except Exception as e:
            logger.warning(f"_format_display_formula: 主方案失败: {e}，使用降级模式")
            self._format_display_formula_fallback(para, equation_number, bookmark_name, bookmark_id)

    def _remove_existing_formula_number(self, para, logger) -> None:
        """移除原始文档中已有的公式编号（避免双编号）

        支持多种编号格式：
        - 英文半角：(1-1), (2-3)
        - 中文全角：（1-1），（2-3）
        - 带空格： (1-1) ， (1-1)

        只处理纯文本 run，保留 OMML 数学公式内容。

        Args:
            para: 段落对象
            logger: 日志记录器
        """
        import re

        patterns = [
            re.compile(r"^[\s\(（]*\d+[-‐–—]\d+[\s\)）]*$"),
            re.compile(r"^[\s\(（]*\d+[-‐–—]\d+[\s\)）]*$"),
            re.compile(r"[\(（]\d+[-‐–—]\d+[\)）]"),
        ]

        runs_to_remove = []

        for i, run in enumerate(para.runs):
            if self._run_contains_math(run):
                continue

            text = run.text
            if not text or not text.strip():
                continue

            stripped = text.strip()

            for pattern in patterns:
                if pattern.match(stripped) or pattern.match(text):
                    runs_to_remove.append(i)
                    logger.debug(f"  检测到已有编号文本，将移除: '{text}'")
                    break

        for i in reversed(runs_to_remove):
            try:
                run_element = para.runs[i]._r
                para._p.remove(run_element)
                logger.debug(f"  已移除第 {i} 个 run 中的编号")
            except Exception as e:
                logger.debug(f"  移除编号失败: {e}")

        for i, run in enumerate(list(para.runs)):
            if self._run_contains_math(run):
                continue

            text = run.text
            if not text or not text.strip():
                continue

            stripped = text.strip()

            for pattern in patterns[2:]:
                match = pattern.search(stripped)
                if match:
                    before = stripped[: match.start()].strip()
                    after = stripped[match.end() :].strip()

                    if len(before) == 0 and len(after) == 0:
                        try:
                            run_element = run._r
                            para._p.remove(run_element)
                            logger.debug(f"  检测到纯编号 run 并移除: '{text}'")
                        except Exception:
                            pass
                    elif len(before) == 0 and len(after) > 0:
                        try:
                            run.text = after
                            logger.debug(f"  移除编号前缀，保留后缀: '{after}'")
                        except Exception:
                            pass
                    elif len(before) > 0 and len(after) == 0:
                        try:
                            run.text = before
                            logger.debug(f"  移除编号后缀，保留前缀: '{before}'")
                        except Exception:
                            pass

    def _format_display_formula_safe(
        self, para, equation_number: str, bookmark_name: str, bookmark_id: str, logger
    ) -> None:
        """安全的公式格式化方法：最小化修改策略

        核心原则：
        1. 不删除原有 OMML 内容
        2. 在段落开头插入书签起始和居中制表符
        3. 在段落末尾插入右对齐制表符、编号和书签结束
        4. 保留 OMML 命名空间完整性

        最终段落结构：
        [bookmarkStart][tab居中][OMML公式...][tab右对齐][(1-1)][bookmarkEnd]

        Args:
            para: 段落对象
            equation_number: 公式编号字符串，如 "(1-1)"
            bookmark_name: 书签名称
            bookmark_id: 书签ID
            logger: 日志记录器
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        try:
            bookmark_start = OxmlElement("w:bookmarkStart")
            bookmark_start.set(qn("w:id"), bookmark_id)
            bookmark_start.set(qn("w:name"), bookmark_name)

            para._p.insert(0, bookmark_start)
            logger.debug(f"  插入书签起始: id={bookmark_id}, name={bookmark_name}")

            tab_run_center = para.add_run("\t")
            tab_run_center.font.name = self.FONT_SETTINGS["english_font"]
            tab_run_center.font.size = Pt(self.FONT_SIZES["body"])

            tab_center_elem = tab_run_center._r
            para._p.insert(1, tab_center_elem)
            logger.debug("  在段首插入居中制表符")

            tab_run_right = para.add_run("\t")
            tab_run_right.font.name = self.FONT_SETTINGS["english_font"]
            tab_run_right.font.size = Pt(self.FONT_SIZES["body"])

            num_run = para.add_run(equation_number)
            num_run.font.name = self.FONT_SETTINGS["english_font"]
            num_run.font.size = Pt(self.FONT_SIZES["body"])
            num_run.font.color.rgb = RGBColor(0, 0, 0)

            bookmark_end = OxmlElement("w:bookmarkEnd")
            bookmark_end.set(qn("w:id"), bookmark_id)
            para._p.append(bookmark_end)
            logger.debug(f"  在段尾添加右制表符+编号+书签结束: {equation_number}")

            logger.debug("_format_display_formula_safe: 成功完成公式格式化")

        except Exception as e:
            logger.warning(f"_format_display_formula_safe: 安全方案失败: {e}")
            raise

    def _format_display_formula_fallback(
        self, para, equation_number: str, bookmark_name: str, bookmark_id: str
    ) -> None:
        """降级方案：确保公式基本可用性

        当主方案失败时，采用最简单的格式化方式：
        - 设置段落居中对齐
        - 统一字体样式
        - 在段落末尾添加编号

        Args:
            para: 段落对象
            equation_number: 公式编号字符串
            bookmark_name: 书签名称
            bookmark_id: 书签ID
        """
        import logging

        logger = logging.getLogger(__name__)

        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt, RGBColor
        except ImportError:
            return

        logger.warning("_format_display_formula_fallback: 使用降级模式处理公式")

        try:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = None

            for run in para.runs:
                run.font.name = self.FONT_SETTINGS["english_font"]
                run.font.size = Pt(self.FONT_SIZES["body"])
                run.font.color.rgb = RGBColor(0, 0, 0)

            num_run = para.add_run("\t" + equation_number)
            num_run.font.name = self.FONT_SETTINGS["english_font"]
            num_run.font.size = Pt(self.FONT_SIZES["body"])
            num_run.font.color.rgb = RGBColor(0, 0, 0)

            logger.debug("  降级模式完成：居中对齐 + 末尾编号")

        except Exception as e:
            logger.error(f"_format_display_formula_fallback: 降级方案也失败: {e}")

    def format_code_blocks(self, doc) -> None:
        """格式化代码块段落（Source Code 样式或等宽字体段落）

        - 使用 Consolas/Courier New 等宽字体
        - 字号小五号(9pt)
        - 段落首行缩进取消
        - 可选浅灰背景

        Args:
            doc: 文档对象
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt, RGBColor
        except ImportError:
            return

        code_style_names = {
            "source code",
            "code",
            "sourcecode",
            "preformatted",
            "verbatim",
            "codesnippet",
            "CodeSnippet",
        }

        for para in doc.paragraphs:
            is_code_block = False

            style_name = ""
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if any(cs in style_name for cs in code_style_names):
                    is_code_block = True

            if not is_code_block:
                text = para.text
                if text and len(text) > 3:
                    lines = text.split("\n")
                    if len(lines) >= 2:
                        all_indented = all(
                            line.startswith(
                                ("    ", "\t", ">>> ", "... ", "# ", "// ", "/* ", "* ")
                            )
                            for line in lines
                            if line.strip()
                        )
                        if all_indented and any(
                            kw in text.lower()
                            for kw in [
                                "def ",
                                "class ",
                                "import ",
                                "return ",
                                "if ",
                                "for ",
                                "function ",
                                "var ",
                                "const ",
                                "=>",
                                "{",
                                "}",
                                ";",
                            ]
                        ):
                            is_code_block = True

            if not is_code_block:
                continue

            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.left_indent = Cm(0.5)
            para.paragraph_format.right_indent = Cm(0.5)

            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F5F5F5")

            pPr = para._p.get_or_add_pPr()
            existing_shd = pPr.find(qn("w:shd"))
            if existing_shd is not None:
                pPr.remove(existing_shd)
            pPr.append(shd)

            for run in para.runs:
                # 代码块使用 Times New Roman 字体，小四字号 (12pt)
                run.font.name = self.FONT_SETTINGS["english_font"]
                run.font.size = Pt(self.FONT_SIZES["body"])
                run.font.color.rgb = RGBColor(0, 0, 0)

                rPr = run._r.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is not None:
                    rFonts.set(qn("w:ascii"), self.FONT_SETTINGS["english_font"])
                    rFonts.set(qn("w:hAnsi"), self.FONT_SETTINGS["english_font"])
                    rFonts.set(qn("w:eastAsia"), self.FONT_SETTINGS["chinese_font"])
