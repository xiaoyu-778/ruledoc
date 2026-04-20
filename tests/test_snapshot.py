"""
快照测试
"""

import hashlib
import os

import pytest
from docx import Document

from ruledoc.formatter import Formatter
from ruledoc.rules import load_rule


class TestSnapshot:
    """快照测试"""

    @pytest.fixture
    def snapshot_dir(self, tmp_path):
        """快照目录"""
        return tmp_path / "snapshots"

    @pytest.fixture
    def input_docx(self, tmp_path):
        """创建标准测试文档"""
        doc = Document()

        doc.add_heading("摘要", level=1)
        doc.add_paragraph("这是摘要的内容，用于测试摘要格式。")

        doc.add_heading("关键词", level=2)
        doc.add_paragraph("关键词1；关键词2；关键词3")

        doc.add_heading("第一章 绪论", level=1)
        doc.add_paragraph("这是第一章的内容，用于测试正文格式。参考文献引用示例[1]。")

        doc.add_heading("1.1 研究背景", level=2)
        doc.add_paragraph("这是研究背景的内容。")

        doc.add_heading("1.2 研究目的", level=2)
        doc.add_paragraph("这是研究目的的内容。")

        doc.add_heading("第二章 方法", level=1)
        doc.add_paragraph("这是第二章的内容。")

        p_fig = doc.add_paragraph("图 2-1 流程图")
        p_fig.style = "Caption"

        p_tab = doc.add_paragraph("表 2-1 数据表")
        p_tab.style = "Caption"

        doc.add_heading("参考文献", level=1)
        doc.add_paragraph("[1] 张三. 测试文献一. 测试期刊, 2023.")
        doc.add_paragraph("[2] 李四. 测试文献二. 测试期刊, 2022.")

        doc.add_heading("致谢", level=1)
        doc.add_paragraph("感谢导师的指导。")

        file_path = tmp_path / "snapshot_input.docx"
        doc.save(str(file_path))

        return str(file_path)

    def _compute_file_hash(self, file_path: str) -> str:
        """计算文件哈希值"""
        hasher = hashlib.sha256()
        with open(file_path, "rb") as f:
            hasher.update(f.read())
        return hasher.hexdigest()

    def test_yzu_snapshot(self, input_docx, tmp_path):
        """YZU 规则快照测试"""
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "yzu_output.docx")

        formatter = Formatter(input_docx, rule, output_path)
        formatter.process()

        assert os.path.exists(output_path)

        doc = Document(output_path)

        assert len(doc.sections) >= 1
        section = doc.sections[0]
        assert section.top_margin is not None
        assert abs(section.top_margin.cm - 2.2) < 0.1

        heading_found = False
        for para in doc.paragraphs:
            style = para.style
            if style is not None and style.name is not None and style.name.startswith("Heading"):
                heading_found = True
                break
        assert heading_found

    def test_yzu_design_snapshot(self, input_docx, tmp_path):
        """YZU Design 规则快照测试"""
        rule = load_rule("yzu_design")
        assert rule is not None
        output_path = str(tmp_path / "yzu_design_output.docx")

        formatter = Formatter(input_docx, rule, output_path)
        formatter.process()

        assert os.path.exists(output_path)

        doc = Document(output_path)

        assert len(doc.sections) >= 1

    def test_output_consistency(self, input_docx, tmp_path):
        """测试输出内容一致性（非文件哈希）"""
        rule = load_rule("yzu")
        assert rule is not None

        output1 = str(tmp_path / "output1.docx")
        output2 = str(tmp_path / "output2.docx")

        formatter1 = Formatter(input_docx, rule, output1)
        formatter1.process()

        formatter2 = Formatter(input_docx, rule, output2)
        formatter2.process()

        doc1 = Document(output1)
        doc2 = Document(output2)

        assert len(doc1.paragraphs) == len(doc2.paragraphs), "段落数应一致"

        for p1, p2 in zip(doc1.paragraphs, doc2.paragraphs):
            assert p1.text == p2.text, f"段落内容应一致: '{p1.text[:30]}' vs '{p2.text[:30]}'"

        assert len(doc1.sections) == len(doc2.sections), "章节数应一致"

    def test_paragraph_count_preserved(self, input_docx, tmp_path):
        """测试段落数保持"""
        input_doc = Document(input_docx)
        input_count = len(input_doc.paragraphs)

        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(input_docx, rule, output_path)
        formatter.process()

        output_doc = Document(output_path)
        output_count = len(output_doc.paragraphs)

        assert output_count == input_count, "格式化不应改变段落数"


class TestSnapshotValidation:
    """快照验证测试"""

    def test_page_settings_validation(self, tmp_path):
        """测试页面设置验证"""
        doc = Document()
        doc.add_paragraph("测试")

        input_path = str(tmp_path / "input.docx")
        output_path = str(tmp_path / "output.docx")
        doc.save(input_path)

        rule = load_rule("yzu")
        assert rule is not None
        formatter = Formatter(input_path, rule, output_path)
        formatter.process()

        output_doc = Document(output_path)
        section = output_doc.sections[0]

        settings = rule.get_validated_page_settings()

        assert section.top_margin is not None
        assert section.bottom_margin is not None
        assert section.left_margin is not None
        assert section.right_margin is not None
        assert abs(section.top_margin.cm - settings["top_margin"]) < 0.1
        assert abs(section.bottom_margin.cm - settings["bottom_margin"]) < 0.1
        assert abs(section.left_margin.cm - settings["left_margin"]) < 0.1
        assert abs(section.right_margin.cm - settings["right_margin"]) < 0.1

    def test_heading_styles_preserved(self, tmp_path):
        """测试标题样式保持

        注意：论文标题（第一个 Heading 1）会被移除 Heading 样式，
        因为它不是一级标题，而是独立的论文标题格式。
        所以这里只检查 Heading 2 和 Heading 3 的存在。
        """
        doc = Document()
        doc.add_heading("论文标题", level=1)
        doc.add_heading("第一章 绪论", level=1)
        doc.add_heading("二级标题", level=2)
        doc.add_heading("三级标题", level=3)

        input_path = str(tmp_path / "input.docx")
        output_path = str(tmp_path / "output.docx")
        doc.save(input_path)

        rule = load_rule("yzu")
        assert rule is not None
        formatter = Formatter(input_path, rule, output_path)
        formatter.process()

        output_doc = Document(output_path)

        heading_levels = []
        for para in output_doc.paragraphs:
            style = para.style
            if style is not None and style.name is not None and style.name.startswith("Heading"):
                level = int(style.name.split()[-1])
                heading_levels.append(level)

        assert 1 in heading_levels
        assert 2 in heading_levels
        assert 3 in heading_levels
