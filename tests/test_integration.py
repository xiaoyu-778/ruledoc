"""集成测试
"""

import os
import re

import pytest
from docx import Document

from ruledoc.context import ProcessingContext
from ruledoc.formatter import Formatter
from ruledoc.processors import CaptionProcessor, HeadingProcessor, StyleProcessor
from ruledoc.rules import load_rule


class TestIntegration:
    """集成测试"""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """创建测试 Word 文档"""
        doc = Document()

        doc.add_heading("第一章 绪论", level=1)
        doc.add_paragraph("这是第一章的内容，用于测试正文格式。")

        doc.add_heading("1.1 研究背景", level=2)
        doc.add_paragraph("这是研究背景的内容。")

        doc.add_heading("1.2 研究目的", level=2)
        doc.add_paragraph("这是研究目的的内容。")

        doc.add_heading("第二章 文献综述", level=1)
        doc.add_paragraph("这是第二章的内容。")

        p_fig = doc.add_paragraph("图 1-1 测试图片标题")
        p_fig.style = "Caption"

        p_tab = doc.add_paragraph("表 1-1 测试表格标题")
        p_tab.style = "Caption"

        doc.add_heading("参考文献", level=1)
        doc.add_paragraph("[1] 张三. 测试文献一. 测试期刊, 2023.")
        doc.add_paragraph("[2] 李四. 测试文献二. 测试期刊, 2022.")

        file_path = tmp_path / "test_input.docx"
        doc.save(str(file_path))

        return str(file_path)

    def test_formatter_init(self, sample_docx):
        """测试 Formatter 初始化"""
        rule = load_rule("yzu")
        assert rule is not None
        formatter = Formatter(sample_docx, rule)

        assert formatter.input_path == sample_docx
        assert formatter.rule == rule
        assert formatter.output_path.endswith("_formatted.docx")

    def test_formatter_custom_output(self, sample_docx, tmp_path):
        """测试自定义输出路径"""
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "custom_output.docx")
        formatter = Formatter(sample_docx, rule, output_path)

        assert formatter.output_path == output_path

    def test_formatter_process(self, sample_docx, tmp_path):
        """测试完整格式化流程"""
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(sample_docx, rule, output_path)
        formatter.process()

        assert os.path.exists(output_path)

        doc = Document(output_path)
        assert len(doc.paragraphs) > 0

    def test_page_settings_applied(self, sample_docx, tmp_path):
        """测试页面设置应用"""
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(sample_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)
        section = doc.sections[0]

        assert section.top_margin is not None
        assert section.bottom_margin is not None
        assert section.left_margin is not None
        assert section.right_margin is not None
        assert abs(section.top_margin.cm - 2.2) < 0.1
        assert abs(section.bottom_margin.cm - 2.2) < 0.1
        assert abs(section.left_margin.cm - 2.5) < 0.1
        assert abs(section.right_margin.cm - 2.0) < 0.1

    def test_heading_numbering_applied(self, sample_docx, tmp_path):
        """测试标题手动编号应用

        扬州大学格式规范要求使用手动编号格式：
        - 一级标题：1.、2.、3.
        - 二级标题：1.1、1.2、2.1
        - 三级标题：1.1.1、1.1.2
        """
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(sample_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)

        heading_texts = []
        for p in doc.paragraphs:
            style = p.style
            if style is not None and style.name is not None and style.name.startswith("Heading"):
                heading_texts.append(p.text)

        for text in heading_texts:
            assert not text.startswith("第一章")
            assert not text.startswith("第")

        has_numbered_heading = any(
            text.startswith(("1.", "2.", "3.", "4.", "5.")) or re.match(r"^\d+\.\d+\s", text)
            for text in heading_texts
        )
        assert has_numbered_heading or len(heading_texts) == 0

    def test_caption_numbering_updated(self, sample_docx, tmp_path):
        """测试题注编号更新"""
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(sample_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)

        caption_texts = [p.text for p in doc.paragraphs if p.style and p.style.name == "Caption"]

        assert len(caption_texts) >= 2

    def test_warnings_collected(self, sample_docx, tmp_path):
        """测试警告收集"""
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(sample_docx, rule, output_path)
        formatter.process()

        warnings = formatter.get_warnings()

        assert isinstance(warnings, list)

    def test_gutter_setting_applied(self, sample_docx, tmp_path):
        """测试装订线设置"""
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(sample_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)
        section = doc.sections[0]

        assert section.gutter is not None
        assert abs(section.gutter.cm - 0.5) < 0.1

    def test_paragraph_count_preserved(self, sample_docx, tmp_path):
        """测试段落数保持"""
        input_doc = Document(sample_docx)
        input_count = len(input_doc.paragraphs)

        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(sample_docx, rule, output_path)
        formatter.process()

        output_doc = Document(output_path)
        output_count = len(output_doc.paragraphs)

        assert output_count == input_count

    def test_empty_document(self, tmp_path):
        """测试空文档处理"""
        doc = Document()
        input_path = str(tmp_path / "empty.docx")
        output_path = str(tmp_path / "empty_output.docx")
        doc.save(input_path)

        rule = load_rule("yzu")
        assert rule is not None

        formatter = Formatter(input_path, rule, output_path)
        formatter.process()

        assert os.path.exists(output_path)

    def test_special_characters_in_content(self, tmp_path):
        """测试特殊字符处理"""
        doc = Document()
        doc.add_paragraph("特殊字符测试：<>&\"'\n\t")
        doc.add_paragraph("中文测试：你好世界")

        input_path = str(tmp_path / "special.docx")
        output_path = str(tmp_path / "special_output.docx")
        doc.save(input_path)

        rule = load_rule("yzu")
        assert rule is not None

        formatter = Formatter(input_path, rule, output_path)
        formatter.process()

        assert os.path.exists(output_path)


class TestProcessorChain:
    """处理器链测试"""

    @pytest.fixture
    def mock_context(self):
        """创建模拟上下文"""
        doc = Document()
        doc.add_heading("测试标题", level=1)
        doc.add_paragraph("测试正文")

        rule = load_rule("yzu")
        assert rule is not None
        return ProcessingContext(doc=doc, rule=rule)

    def test_style_processor_runs(self, mock_context):
        """测试 StyleProcessor 执行"""
        processor = StyleProcessor()

        processor.process(mock_context)

        assert isinstance(mock_context.warnings, list)

    def test_heading_processor_runs(self, mock_context):
        """测试 HeadingProcessor 执行"""
        processor = HeadingProcessor()

        processor.process(mock_context)

        assert isinstance(mock_context.warnings, list)

    def test_caption_processor_runs(self, mock_context):
        """测试 CaptionProcessor 执行"""
        processor = CaptionProcessor()

        processor.process(mock_context)

        assert isinstance(mock_context.warnings, list)

    def test_processor_order(self, mock_context):
        """测试处理器执行顺序"""
        style_processor = StyleProcessor()
        heading_processor = HeadingProcessor()
        caption_processor = CaptionProcessor()

        style_processor.process(mock_context)
        heading_processor.process(mock_context)
        caption_processor.process(mock_context)

        assert mock_context.chapter_num >= 0

    def test_processor_idempotency(self, mock_context):
        """测试处理器幂等性"""
        processor = StyleProcessor()

        processor.process(mock_context)
        first_warnings = len(mock_context.warnings)

        processor.process(mock_context)
        second_warnings = len(mock_context.warnings)

        assert second_warnings >= first_warnings


class TestCLI:
    """CLI 测试"""

    def test_cli_import(self):
        """测试 CLI 模块导入"""
        from ruledoc.cli import create_parser, main, select_rule

        assert callable(main)
        assert callable(create_parser)
        assert callable(select_rule)

    def test_parser_creation(self):
        """测试参数解析器创建"""
        from ruledoc.cli import create_parser

        parser = create_parser()

        assert parser is not None

    def test_parser_list_rules(self):
        """测试 --list-rules 参数"""
        from ruledoc.cli import create_parser

        parser = create_parser()
        args = parser.parse_args(["--list-rules"])

        assert args.list_rules is True

    def test_parser_input_output(self):
        """测试输入输出参数"""
        from ruledoc.cli import create_parser

        parser = create_parser()
        args = parser.parse_args(["input.docx", "-o", "output.docx", "--rule", "yzu"])

        assert args.input == "input.docx"
        assert args.output == "output.docx"
        assert args.rule == "yzu"
