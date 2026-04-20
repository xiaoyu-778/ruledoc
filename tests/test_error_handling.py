"""异常处理测试

覆盖:
- PandocNotInstalledError
- ProcessingError
- ConfigurationError
- 文件不存在
- 无效输入
- 边界条件
"""

import os
from unittest.mock import patch

import pytest
from docx import Document

from ruledoc.exceptions import (
    ConfigurationError,
    PandocNotInstalledError,
    ProcessingError,
    RuleDocError,
    RuleNotFoundError,
)
from ruledoc.formatter import Formatter
from ruledoc.pandoc_converter import PandocConverter
from ruledoc.rules import load_rule


class TestPandocNotInstalledError:
    """Pandoc 未安装异常测试"""

    @patch.object(PandocConverter, "is_available")
    def test_pandoc_not_installed_raises_error(self, mock_is_available, tmp_path):
        """测试 Pandoc 未安装时抛出异常"""
        mock_is_available.return_value = False
        PandocConverter._available = False

        md_file = tmp_path / "test.md"
        md_file.write_text("# Test")

        converter = PandocConverter()

        with pytest.raises(PandocNotInstalledError) as exc_info:
            converter.convert(str(md_file))

        assert "Pandoc" in str(exc_info.value)
        assert exc_info.value.context.get("download_url") is not None

    def test_pandoc_not_installed_error_message(self):
        """测试 Pandoc 未安装错误消息"""
        error = PandocNotInstalledError(
            "Pandoc 未安装", {"download_url": "https://pandoc.org/installing.html"}
        )

        assert "Pandoc" in error.message
        assert error.context["download_url"] == "https://pandoc.org/installing.html"


class TestFileNotFoundError:
    """文件不存在异常测试"""

    def test_input_file_not_found(self):
        """测试输入文件不存在"""
        rule = load_rule("yzu")
        assert rule is not None

        with pytest.raises(ProcessingError) as exc_info:
            Formatter("/nonexistent/path/file.docx", rule)

        assert "不存在" in str(exc_info.value)

    def test_input_file_not_found_md(self):
        """测试 Markdown 文件不存在"""
        rule = load_rule("yzu")
        assert rule is not None

        with pytest.raises(ProcessingError) as exc_info:
            Formatter("/nonexistent/path/file.md", rule)

        assert "不存在" in str(exc_info.value)


class TestInvalidInputError:
    """无效输入异常测试"""

    def test_unsupported_format(self, tmp_path):
        """测试不支持的文件格式"""
        rule = load_rule("yzu")
        assert rule is not None

        invalid_file = tmp_path / "test.xyz"
        invalid_file.write_text("test content")

        with pytest.raises(ConfigurationError) as exc_info:
            Formatter(str(invalid_file), rule)

        assert "不支持" in str(exc_info.value)

    def test_file_too_large(self, tmp_path):
        """测试文件过大"""
        rule = load_rule("yzu")
        assert rule is not None

        large_file = tmp_path / "large.docx"

        with patch("os.path.getsize", return_value=200 * 1024 * 1024):
            with patch("os.path.exists", return_value=True):
                with pytest.raises(ConfigurationError) as exc_info:
                    Formatter(str(large_file), rule)

                assert "过大" in str(exc_info.value) or "超过限制" in str(exc_info.value)


class TestProcessingError:
    """处理错误异常测试"""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """创建测试文档"""
        doc = Document()
        doc.add_heading("测试", level=1)
        doc.add_paragraph("内容")

        file_path = tmp_path / "test.docx"
        doc.save(str(file_path))

        return str(file_path)

    def test_processing_error_with_context(self):
        """测试带上下文的处理错误"""
        error = ProcessingError("文档处理失败", {"input_path": "test.docx", "error": "格式错误"})

        assert error.message == "文档处理失败"
        assert error.context["input_path"] == "test.docx"

    def test_formatter_save_error(self, sample_docx, tmp_path):
        """测试保存错误"""
        rule = load_rule("yzu")
        assert rule is not None

        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(sample_docx, rule, output_path)

        with patch("builtins.open", side_effect=PermissionError("Permission denied")):
            with pytest.raises(ProcessingError):
                formatter.process()


class TestRuleNotFoundError:
    """规则不存在异常测试"""

    def test_rule_not_found_error(self):
        """测试规则不存在错误"""
        rule = load_rule("nonexistent_rule_xyz")

        assert rule is None

    def test_rule_not_found_error_message(self):
        """测试规则不存在错误消息"""
        error = RuleNotFoundError(
            "规则 'xyz' 不存在", {"available_rules": ["yzu_thesis", "yzu_design"]}
        )

        assert "xyz" in error.message
        assert error.context["available_rules"] == ["yzu_thesis", "yzu_design"]


class TestConfigurationError:
    """配置错误异常测试"""

    def test_configuration_error(self):
        """测试配置错误"""
        error = ConfigurationError("无效的配置", {"key": "top_margin", "value": -1})

        assert "无效" in error.message
        assert error.context["key"] == "top_margin"


class TestExceptionInheritance:
    """异常继承关系测试"""

    def test_all_exceptions_inherit_from_base(self):
        """测试所有异常继承自基类"""
        exceptions = [
            RuleNotFoundError("test"),
            ProcessingError("test"),
            ConfigurationError("test"),
            PandocNotInstalledError("test"),
        ]

        for exc in exceptions:
            assert isinstance(exc, RuleDocError)
            assert isinstance(exc, Exception)

    def test_catch_with_base_class(self):
        """测试使用基类捕获异常"""
        with pytest.raises(RuleDocError):
            raise RuleNotFoundError("test")

        with pytest.raises(RuleDocError):
            raise ProcessingError("test")

        with pytest.raises(RuleDocError):
            raise ConfigurationError("test")

        with pytest.raises(RuleDocError):
            raise PandocNotInstalledError("test")


class TestEdgeCases:
    """边界条件测试"""

    @pytest.fixture
    def empty_docx(self, tmp_path):
        """创建空文档"""
        doc = Document()

        file_path = tmp_path / "empty.docx"
        doc.save(str(file_path))

        return str(file_path)

    @pytest.fixture
    def special_chars_docx(self, tmp_path):
        """创建包含特殊字符的文档"""
        doc = Document()
        doc.add_heading("测试特殊字符 <>&\"'", level=1)
        doc.add_paragraph("包含特殊字符: < > & \" ' \\ / \n \t")
        doc.add_paragraph("Unicode: 中文 日本語 한국어 العربية")
        doc.add_paragraph("Emoji: 😀 🎉 📝 ✅")

        file_path = tmp_path / "special_chars.docx"
        doc.save(str(file_path))

        return str(file_path)

    def test_empty_document(self, empty_docx, tmp_path):
        """测试空文档处理"""
        rule = load_rule("yzu")
        assert rule is not None

        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(empty_docx, rule, output_path)
        formatter.process()

        assert os.path.exists(output_path)

        doc = Document(output_path)
        assert len(doc.paragraphs) >= 0

    def test_special_characters(self, special_chars_docx, tmp_path):
        """测试特殊字符处理"""
        rule = load_rule("yzu")
        assert rule is not None

        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(special_chars_docx, rule, output_path)
        formatter.process()

        assert os.path.exists(output_path)

        doc = Document(output_path)
        assert len(doc.paragraphs) > 0

    def test_long_paragraph(self, tmp_path):
        """测试超长段落"""
        doc = Document()

        long_text = "测试内容 " * 10000
        doc.add_paragraph(long_text)

        file_path = tmp_path / "long_para.docx"
        doc.save(str(file_path))

        rule = load_rule("yzu")
        assert rule is not None

        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(str(file_path), rule, output_path)
        formatter.process()

        assert os.path.exists(output_path)

    def test_many_paragraphs(self, tmp_path):
        """测试大量段落"""
        doc = Document()

        for i in range(1000):
            doc.add_paragraph(f"段落 {i}: 这是测试内容。")

        file_path = tmp_path / "many_paras.docx"
        doc.save(str(file_path))

        rule = load_rule("yzu")
        assert rule is not None

        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(str(file_path), rule, output_path)
        formatter.process()

        assert os.path.exists(output_path)

        output_doc = Document(output_path)
        assert len(output_doc.paragraphs) == 1000

    def test_deeply_nested_headings(self, tmp_path):
        """测试深层嵌套标题"""
        doc = Document()

        for level in range(1, 7):
            doc.add_heading(f"标题 Level {level}", level=level)
            doc.add_paragraph(f"Level {level} 的内容")

        file_path = tmp_path / "nested.docx"
        doc.save(str(file_path))

        rule = load_rule("yzu")
        assert rule is not None

        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(str(file_path), rule, output_path)
        formatter.process()

        assert os.path.exists(output_path)

    def test_mixed_content(self, tmp_path):
        """测试混合内容"""
        doc = Document()

        doc.add_heading("第一章", level=1)
        doc.add_paragraph("正文内容")

        doc.add_heading("1.1 小节", level=2)
        doc.add_paragraph("图 1-1 测试图片")
        doc.add_paragraph("表 1-1 测试表格")

        doc.add_heading("参考文献", level=1)
        doc.add_paragraph("[1] 文献一")
        doc.add_paragraph("[2] 文献二")

        doc.add_heading("致谢", level=1)
        doc.add_paragraph("感谢...")

        file_path = tmp_path / "mixed.docx"
        doc.save(str(file_path))

        rule = load_rule("yzu")
        assert rule is not None

        output_path = str(tmp_path / "output.docx")

        formatter = Formatter(str(file_path), rule, output_path)
        formatter.process()

        assert os.path.exists(output_path)


class TestErrorRecovery:
    """错误恢复测试"""

    def test_temp_file_cleanup_on_error(self, tmp_path):
        """测试错误时临时文件清理"""
        rule = load_rule("yzu")
        assert rule is not None

        md_file = tmp_path / "test.md"
        md_file.write_text("# Test")

        with patch("ruledoc.pandoc_converter.PandocConverter.is_available", return_value=False):
            PandocConverter._available = False

            try:
                formatter = Formatter(str(md_file), rule)
                formatter.process()
            except (ProcessingError, PandocNotInstalledError):
                pass

        temp_files = list(tmp_path.glob("*.docx"))
        for f in temp_files:
            if "test_formatted" not in str(f):
                assert not str(f).endswith(".docx") or "test" not in str(f)
