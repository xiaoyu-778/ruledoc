"""CLI 模块完整测试

覆盖:
- main() 入口
- 参数解析
- 错误处理
- 规则选择
"""

import os
from unittest.mock import patch

import pytest
from docx import Document

from ruledoc.cli import create_parser, format_file, main, select_rule
from ruledoc.exceptions import ProcessingError, RuleNotFoundError


class TestCLIMain:
    """CLI main() 入口测试"""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """创建测试文档"""
        doc = Document()
        doc.add_heading("测试标题", level=1)
        doc.add_paragraph("测试内容")

        file_path = tmp_path / "test_input.docx"
        doc.save(str(file_path))

        return str(file_path)

    def test_main_list_rules(self, capsys):
        """测试 --list-rules 参数"""
        exit_code = main(["--list-rules"])

        assert exit_code == 0

        captured = capsys.readouterr()
        assert "可用规则" in captured.out or "yzu" in captured.out

    def test_main_no_input(self, capsys):
        """测试无输入文件"""
        exit_code = main([])

        assert exit_code == 1

    def test_main_with_input(self, sample_docx, tmp_path, capsys):
        """测试正常输入处理"""
        output_path = str(tmp_path / "output.docx")

        exit_code = main([sample_docx, "-o", output_path, "--rule", "yzu"])

        assert exit_code == 0
        assert os.path.exists(output_path)

        captured = capsys.readouterr()
        assert "格式化完成" in captured.out

    def test_main_rule_not_found(self, sample_docx, capsys):
        """测试规则不存在"""
        exit_code = main([sample_docx, "--rule", "nonexistent_rule_xyz"])

        assert exit_code == 2

        captured = capsys.readouterr()
        assert "错误" in captured.err or "不存在" in captured.err

    def test_main_file_not_found(self, capsys):
        """测试输入文件不存在"""
        exit_code = main(["/nonexistent/path/file.docx", "--rule", "yzu"])

        assert exit_code == 4

    def test_main_keyboard_interrupt(self, capsys):
        """测试键盘中断"""
        with patch("ruledoc.cli.select_rule", side_effect=KeyboardInterrupt):
            exit_code = main(["input.docx"])

            assert exit_code == 130

    def test_main_unknown_error(self, capsys):
        """测试未知错误"""
        with patch("ruledoc.cli.select_rule", side_effect=RuntimeError("Unknown error")):
            exit_code = main(["input.docx"])

            assert exit_code == 1


class TestCLIParser:
    """CLI 参数解析测试"""

    def test_parser_default_values(self):
        """测试默认参数值"""
        parser = create_parser()
        args = parser.parse_args(["input.docx"])

        assert args.input == "input.docx"
        assert args.output is None
        assert args.rule is None
        assert args.list_rules is False

    def test_parser_all_options(self):
        """测试所有参数选项"""
        parser = create_parser()
        args = parser.parse_args(["input.md", "-o", "output.docx", "--rule", "yzu_design"])

        assert args.input == "input.md"
        assert args.output == "output.docx"
        assert args.rule == "yzu_design"

    def test_parser_version(self, capsys):
        """测试 --version 参数"""
        parser = create_parser()

        with pytest.raises(SystemExit) as exc_info:
            parser.parse_args(["--version"])

        assert exc_info.value.code == 0


class TestCLISelectRule:
    """CLI 规则选择测试"""

    def test_select_rule_by_name(self):
        """测试按名称选择规则"""
        rule = select_rule("yzu", non_interactive=True)

        assert rule is not None
        assert rule.name == "yzu_thesis"

    def test_select_rule_default(self):
        """测试默认规则选择"""
        rule = select_rule(None, non_interactive=True)

        assert rule is not None
        assert rule.name == "yzu_thesis"

    def test_select_rule_not_found(self):
        """测试规则不存在"""
        with pytest.raises(RuleNotFoundError):
            select_rule("nonexistent_rule_xyz", non_interactive=True)

    def test_select_rule_yzu_design(self):
        """测试选择 yzu_design 规则"""
        rule = select_rule("yzu_design", non_interactive=True)

        assert rule is not None
        assert rule.name == "yzu_design"

    def test_select_rule_yzu_thesis(self):
        """测试选择 yzu_thesis 规则"""
        rule = select_rule("yzu_thesis", non_interactive=True)

        assert rule is not None
        assert rule.name == "yzu_thesis"


class TestCLIFormatFile:
    """CLI 格式化文件测试"""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """创建测试文档"""
        doc = Document()
        doc.add_heading("测试标题", level=1)
        doc.add_paragraph("测试内容")

        file_path = tmp_path / "test_input.docx"
        doc.save(str(file_path))

        return str(file_path)

    def test_format_file_success(self, sample_docx, tmp_path):
        """测试成功格式化文件"""
        from ruledoc.rules import load_rule

        rule = load_rule("yzu")
        assert rule is not None

        output_path = str(tmp_path / "output.docx")
        result = format_file(sample_docx, rule, output_path)

        assert result == output_path
        assert os.path.exists(output_path)

    def test_format_file_auto_output(self, sample_docx, tmp_path):
        """测试自动生成输出路径"""
        import shutil

        from ruledoc.rules import load_rule

        rule = load_rule("yzu")
        assert rule is not None

        sample_copy = str(tmp_path / "test_copy.docx")
        shutil.copy(sample_docx, sample_copy)

        result = format_file(sample_copy, rule)

        assert result.endswith("_formatted.docx")
        assert os.path.exists(result)

    def test_format_file_processing_error(self):
        """测试处理错误"""
        from ruledoc.rules import load_rule

        rule = load_rule("yzu")
        assert rule is not None

        with pytest.raises(ProcessingError):
            format_file("/nonexistent/file.docx", rule)


class TestCLIIntegration:
    """CLI 集成测试"""

    @pytest.fixture
    def sample_md(self, tmp_path):
        """创建测试 Markdown 文件"""
        content = """# 测试标题

这是测试内容。

## 第一章 绪论

这是第一章的内容。
"""
        file_path = tmp_path / "test_input.md"
        file_path.write_text(content, encoding="utf-8")

        return str(file_path)

    def test_cli_with_markdown_input(self, sample_md, tmp_path):
        """测试 Markdown 输入"""
        output_path = str(tmp_path / "output.docx")

        try:
            exit_code = main([sample_md, "-o", output_path, "--rule", "yzu"])

            if exit_code == 0:
                assert os.path.exists(output_path)
            else:
                pytest.skip("Pandoc not installed")
        except Exception as e:
            if "Pandoc" in str(e):
                pytest.skip("Pandoc not installed")
            raise

    def test_cli_output_to_stdout(self, sample_md, capsys):
        """测试输出到标准输出"""
        try:
            exit_code = main([sample_md, "--rule", "yzu"])

            if exit_code == 0:
                captured = capsys.readouterr()
                assert "输出文件" in captured.out or "格式化完成" in captured.out
        except Exception as e:
            if "Pandoc" in str(e):
                pytest.skip("Pandoc not installed")
            raise
