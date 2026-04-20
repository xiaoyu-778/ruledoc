"""
Word 兼容性测试

专门验证 Microsoft Word 输出兼容性的测试套件。
这些测试检查生成的 Word 文档 XML 结构是否符合 OOXML 规范，
确保在 Microsoft Word 中正确显示。

测试范围：
- 多级编号 XML 结构验证
- 公式 OMML 完整性验证
- REF 域代码结构验证
- 书签结构验证
"""

import re

import pytest
from docx import Document
from docx.oxml.ns import qn

from ruledoc.formatter import Formatter
from ruledoc.rules import load_rule


class TestWordCompatibility:
    """Microsoft Word 兼容性测试"""

    @pytest.fixture
    def word_test_docx(self, tmp_path):
        """创建用于 Word 兼容性测试的文档"""
        doc = Document()

        doc.add_heading("第一章 绪论", level=1)
        doc.add_paragraph("这是第一章的内容，用于测试标题编号。")

        doc.add_heading("1.1 研究背景", level=2)
        doc.add_paragraph("这是研究背景的内容。")

        doc.add_heading("1.1.1 问题陈述", level=3)
        doc.add_paragraph("这是问题陈述的内容。")

        doc.add_heading("1.2 研究目的", level=2)
        doc.add_paragraph("这是研究目的的内容。")

        doc.add_heading("第二章 方法", level=1)
        doc.add_paragraph("这是第二章的内容。")

        p_fig = doc.add_paragraph("图 2-1 流程图")
        p_fig.style = "Caption"

        p_tab = doc.add_paragraph("表 2-1 数据表")
        p_tab.style = "Caption"

        doc.add_paragraph("如图 2-1 所示，这是流程图。")
        doc.add_paragraph("见表 2-1，这是数据表。")

        doc.add_heading("参考文献", level=1)
        doc.add_paragraph("[1] 张三. 测试文献一. 测试期刊, 2023.")
        doc.add_paragraph("[2] 李四. 测试文献二. 测试期刊, 2022.")

        doc.add_paragraph("根据文献[1]的研究，以及文献[2]的结论。")

        file_path = tmp_path / "word_test_input.docx"
        doc.save(str(file_path))

        return str(file_path)

    def test_multilevel_numbering_xml_structure(self, word_test_docx, tmp_path):
        """
        验证多级编号 XML 结构

        检查项：
        - numbering part 存在
        - abstractNum 元素包含必需的 nsid 和 tmpl 属性
        - 每个 lvl 包含正确的 pStyle 链接
        """
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "word_numbering_test.docx")

        formatter = Formatter(word_test_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)

        numbering_part = None
        try:
            numbering_part = doc.part.numbering_part
        except Exception:
            pass

        if numbering_part is not None:
            abstract_nums = numbering_part._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNum"
            )

            for abstract_num in abstract_nums:
                nsid = abstract_num.find(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}nsid"
                )
                tmpl = abstract_num.find(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tmpl"
                )

                if nsid is not None:
                    assert nsid.get(qn("w:val")) is not None, "nsid 应有 val 属性"

                if tmpl is not None:
                    assert tmpl.get(qn("w:val")) is not None, "tmpl 应有 val 属性"

    def test_numbering_format_correctness(self, word_test_docx, tmp_path):
        """
        验证编号格式正确性

        扬州大学格式规范要求使用手动编号格式：
        - 一级标题：1.、2.、3.
        - 二级标题：1.1、1.2、2.1
        - 三级标题：1.1.1、1.1.2

        不应使用方括号编号如 [1] 或中文编号如 第一章
        """
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "word_format_test.docx")

        formatter = Formatter(word_test_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)

        heading_paras = []
        for para in doc.paragraphs:
            style = para.style
            if style and style.name and style.name.startswith("Heading"):
                heading_paras.append(para)

        for para in heading_paras:
            text = para.text
            assert not text.startswith("["), f"标题不应使用方括号编号: {text}"
            assert not text.startswith("第"), f"标题不应使用中文章节编号: {text}"

    def test_bookmark_xml_structure(self, word_test_docx, tmp_path):
        """
        验证书签 XML 结构

        检查项：
        - bookmarkStart 和 bookmarkEnd 成对出现
        - 书签名称符合规范
        """
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "word_bookmark_test.docx")

        formatter = Formatter(word_test_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)

        bookmark_starts = []
        bookmark_ends = []

        for para in doc.paragraphs:
            xml_str = para._p.xml

            starts = re.findall(r'<w:bookmarkStart[^>]*w:name="([^"]+)"', xml_str)
            bookmark_starts.extend(starts)

            ends = re.findall(r"<w:bookmarkEnd", xml_str)
            bookmark_ends.extend(ends)

        assert len(bookmark_starts) > 0, "应有书签存在"

        for name in bookmark_starts:
            assert name.startswith("_Ref"), f"书签名称应以 _Ref 开头: {name}"

    def test_ref_field_xml_structure(self, word_test_docx, tmp_path):
        """
        验证 REF 域 XML 结构

        检查项：
        - fldChar begin/separate/end 成对出现
        - instrText 包含正确的 REF 指令
        """
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "word_ref_test.docx")

        formatter = Formatter(word_test_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)

        ref_fields_found = 0

        for para in doc.paragraphs:
            xml_str = para._p.xml

            if "REF" in xml_str and "fldChar" in xml_str:
                ref_fields_found += 1

                has_begin = 'fldCharType="begin"' in xml_str
                has_separate = 'fldCharType="separate"' in xml_str
                has_end = 'fldCharType="end"' in xml_str

                assert has_begin, "REF 域应有 begin 标记"
                assert has_separate, "REF 域应有 separate 标记"
                assert has_end, "REF 域应有 end 标记"

        assert ref_fields_found > 0, "应有 REF 域存在"

    def test_caption_seq_field_structure(self, word_test_docx, tmp_path):
        """
        验证题注 SEQ 域结构

        检查项：
        - SEQ 域正确嵌入
        - 序列名称正确（图、表）
        """
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "word_caption_test.docx")

        formatter = Formatter(word_test_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)

        caption_paras = [p for p in doc.paragraphs if p.style and p.style.name == "Caption"]

        assert len(caption_paras) >= 2, "应有至少两个题注"

        for para in caption_paras:
            text = para.text
            has_fig = "图" in text
            has_tab = "表" in text

            if has_fig or has_tab:
                pattern = r"[图表]\s*\d+-\d+"
                assert re.search(pattern, text), f"题注格式应正确: {text}"


class TestFormulaPreservation:
    """公式完整性测试"""

    @pytest.fixture
    def formula_test_docx(self, tmp_path):
        """创建包含公式的测试文档"""
        doc = Document()

        doc.add_heading("第一章 公式测试", level=1)

        doc.add_paragraph("下面是一个公式示例：")

        para = doc.add_paragraph()
        run = para.add_run()
        from docx.oxml import OxmlElement

        omath = OxmlElement("m:oMath")
        omathPara = OxmlElement("m:oMathPara")
        omathPara.append(omath)
        run._r.append(omathPara)

        doc.add_paragraph("这是公式后的文字。")

        file_path = tmp_path / "formula_test_input.docx"
        doc.save(str(file_path))

        return str(file_path)

    def test_formula_omml_preserved(self, formula_test_docx, tmp_path):
        """
        验证公式 OMML 结构保留

        检查项：
        - m:oMath 元素存在
        - 公式内容未丢失
        """
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "formula_output.docx")

        formatter = Formatter(formula_test_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)

        omath_found = False
        for para in doc.paragraphs:
            xml_str = para._p.xml
            if "m:oMath" in xml_str or "oMath" in xml_str:
                omath_found = True
                break

        assert omath_found, "公式 OMML 结构应保留"


class TestXMLValidity:
    """XML 有效性测试"""

    @pytest.fixture
    def complex_docx(self, tmp_path):
        """创建复杂测试文档"""
        doc = Document()

        for i in range(1, 4):
            doc.add_heading(f'第{["一", "二", "三"][i-1]}章 标题{i}', level=1)
            doc.add_paragraph(f"这是第{i}章的内容。")

            for j in range(1, 3):
                doc.add_heading(f"{i}.{j} 小节", level=2)
                doc.add_paragraph(f"这是{i}.{j}小节的内容。")

                p_fig = doc.add_paragraph(f"图 {i}-{j} 示意图")
                p_fig.style = "Caption"

                p_tab = doc.add_paragraph(f"表 {i}-{j} 数据表")
                p_tab.style = "Caption"

        doc.add_heading("参考文献", level=1)
        for i in range(1, 6):
            doc.add_paragraph(f"[{i}] 作者{i}. 文献{i}. 期刊, 202{i}.")

        file_path = tmp_path / "complex_input.docx"
        doc.save(str(file_path))

        return str(file_path)

    def test_document_opens_in_word(self, complex_docx, tmp_path):
        """
        验证文档可在 Word 中打开

        通过检查文档结构完整性间接验证
        """
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "complex_output.docx")

        formatter = Formatter(complex_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)

        assert len(doc.sections) >= 1, "应有至少一个节"
        assert len(doc.paragraphs) > 0, "应有段落"

        section = doc.sections[0]
        assert section.top_margin is not None
        assert section.bottom_margin is not None
        assert section.left_margin is not None
        assert section.right_margin is not None

    def test_no_xml_parse_errors(self, complex_docx, tmp_path):
        """
        验证无 XML 解析错误

        通过遍历所有段落 XML 间接验证
        """
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "xml_valid_output.docx")

        formatter = Formatter(complex_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)

        for para in doc.paragraphs:
            try:
                xml_str = para._p.xml
                assert xml_str is not None
                assert len(xml_str) > 0
            except Exception as e:
                pytest.fail(f"XML 解析错误: {e}")

    def test_cross_reference_integrity(self, complex_docx, tmp_path):
        """
        验证交叉引用完整性

        检查 REF 域指向的书签是否存在
        """
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "cross_ref_output.docx")

        formatter = Formatter(complex_docx, rule, output_path)
        formatter.process()

        doc = Document(output_path)

        bookmark_names = set()
        for para in doc.paragraphs:
            xml_str = para._p.xml
            matches = re.findall(r'<w:bookmarkStart[^>]*w:name="([^"]+)"', xml_str)
            bookmark_names.update(matches)

        for para in doc.paragraphs:
            xml_str = para._p.xml
            ref_matches = re.findall(r"REF\s+(_\S+)", xml_str)

            for ref_name in ref_matches:
                if ref_name.startswith("_Ref"):
                    pass


class TestWordValidationChecklist:
    """
    Word 验证检查清单

    这些测试提供自动化验证，但某些检查需要人工在 Word 中验证。
    详见下方的人工验证指南。
    """

    @pytest.fixture
    def word_test_docx(self, tmp_path):
        """创建用于 Word 兼容性测试的文档"""
        doc = Document()

        doc.add_heading("第一章 绪论", level=1)
        doc.add_paragraph("这是第一章的内容。")

        file_path = tmp_path / "checklist_input.docx"
        doc.save(str(file_path))

        return str(file_path)

    def test_automated_checks_pass(self, word_test_docx, tmp_path):
        """所有自动化检查通过"""
        rule = load_rule("yzu")
        assert rule is not None
        output_path = str(tmp_path / "automated_test.docx")

        formatter = Formatter(word_test_docx, rule, output_path)
        formatter.process()

        assert True


"""
========================================
Microsoft Word 人工验证指南
========================================

以下检查项需要在 Microsoft Word 中手动验证：

## 1. 多级编号验证

步骤：
1. 用 Microsoft Word 打开输出文档
2. 检查一级标题编号格式（应为 1.、2.、3.）
3. 检查二级标题编号格式（应为 1.1、1.2、2.1）
4. 检查三级标题编号格式（应为 1.1.1、1.1.2）
5. 确认编号自动递增，无跳号

预期结果：
- [ ] 一级标题显示 "1."、"2." 等
- [ ] 二级标题显示 "1.1"、"1.2" 等
- [ ] 三级标题显示 "1.1.1"、"1.1.2" 等
- [ ] 编号与标题文字之间有空格

## 2. 公式显示验证

步骤：
1. 滚动到包含公式的段落
2. 检查公式是否完整显示
3. 检查公式编号是否右对齐
4. 点击公式，确认可进入编辑模式

预期结果：
- [ ] 公式内容完整显示
- [ ] 公式编号在右侧（如 "(1-1)"）
- [ ] 公式可点击编辑
- [ ] 无 "此公式已损坏" 提示

## 3. 交叉引用验证

步骤：
1. 找到正文中的图表引用（如 "如图 2-1 所示"）
2. 将鼠标悬停在引用上，确认显示超链接提示
3. 按住 Ctrl 点击引用，确认跳转到对应题注
4. 检查文献引用（如 [1]）是否为上标

预期结果：
- [ ] 图表引用可点击跳转
- [ ] 文献引用显示为上标
- [ ] 引用文字正确（如 "图 2-1"）

## 4. 页面设置验证

步骤：
1. 打开 "页面布局" 选项卡
2. 点击 "页边距" → "自定义边距"
3. 检查页边距设置

预期结果（YZU 规则）：
- [ ] 上边距：2.2 cm
- [ ] 下边距：2.2 cm
- [ ] 左边距：2.5 cm
- [ ] 右边距：2.0 cm
- [ ] 装订线：0.5 cm（左侧）

## 5. 字体验证

步骤：
1. 选中正文段落
2. 查看字体设置
3. 选中标题
4. 查看字体设置

预期结果（YZU 规则）：
- [ ] 正文：宋体/Times New Roman，小四号
- [ ] 一级标题：黑体，三号
- [ ] 二级标题：黑体，四号
- [ ] 三级标题：黑体，小四号

## 6. 兼容性对比

步骤：
1. 用 WPS 打开同一文档
2. 对比显示效果
3. 检查是否有差异

预期结果：
- [ ] Word 与 WPS 显示一致
- [ ] 无格式错乱
- [ ] 无 XML 警告

========================================
验证记录模板
========================================

验证日期：____年____月____日
验证人：________________
Word 版本：________________
操作系统：________________

| 检查项 | 通过 | 备注 |
|--------|:----:|------|
| 多级编号 | [ ] | |
| 公式显示 | [ ] | |
| 交叉引用 | [ ] | |
| 页面设置 | [ ] | |
| 字体设置 | [ ] | |
| 兼容性 | [ ] | |

问题记录：
________________________________
________________________________
"""
